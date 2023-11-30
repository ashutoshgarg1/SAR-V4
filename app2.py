#-*- coding: utf-8 -*-
import random,os,json,io,re,zipfile,tempfile
import ssl
import pandas as pd
import streamlit as st
import streamlit_toggle as tog
import cv2
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from langchain import HuggingFaceHub
from langchain.llms import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import openai 
import fitz
import docx
from gtts import gTTS
import PyPDF2
from PyPDF2 import PdfReader
from utils import text_to_docs
from langchain import PromptTemplate, LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from langchain.memory import ConversationSummaryBufferMemory
from langchain.chains.conversation.prompt import ENTITY_MEMORY_CONVERSATION_TEMPLATE
from langchain.chains.conversation.memory import ConversationEntityMemory
from langchain.callbacks import get_openai_callback
from io import StringIO
from io import BytesIO
from usellm import Message, Options, UseLLM
from huggingface_hub import login
import pytesseract
from utils import *
from retr_doc import *
import os
import cv2
import openai
import numpy


# Setting Env
if st.secrets["OPENAI_API_KEY"] is not None:
    os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
else:
    os.environ["OPENAI_API_KEY"] = os.environ.get("OPENAI_API_KEY")

@st.cache_data
def show_pdf(file_path):
    with open(file_path,"rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="1000px" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

@st.cache_data
def pdf_to_bytes(pdf_file_):
    with open(pdf_file_,"rb") as pdf_file:
        pdf_content = pdf_file.read()
        pdf_bytes_io = io.BytesIO(pdf_content)
    return pdf_bytes_io

@st.cache_data
def read_pdf_files(path):
    pdf_files =[]
    directoty_path = path
    files = os.listdir(directoty_path)
    for file in files:
            pdf_files.append(file)
    return pdf_files


@st.cache_data
def merge_pdfs(pdf_list):
    """
    Helper function to merge PDFs
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        pdf_document = PyPDF2.PdfReader(pdf)
        pdf_merger.append(pdf_document)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
  
    return output_pdf


@st.cache_data(show_spinner=False)
def usellm(prompt):
    """
    Getting GPT-3.5 Model into action
    """
    service = UseLLM(service_url="https://usellm.org/api/llm")
    messages = [
      Message(role="system", content="You are a Money Laundering Specialist, who is an expert at finding out suspicious activities"),
      Message(role="user", content=f"{prompt}"),
      ]
    options = Options(messages=messages)
    response = service.chat(options)
    return response.content

# Setting Config for Llama-2
login(token=st.secrets["HUGGINGFACEHUB_API_TOKEN"])
os.environ["HUGGINGFACEHUB_API_TOKEN"] = st.secrets["HUGGINGFACEHUB_API_TOKEN"]


llama_13b = HuggingFaceHub(
            repo_id="meta-llama/Llama-2-13b-chat-hf",
            model_kwargs={"temperature":0.01, 
                        "min_new_tokens":100, 
                        "max_new_tokens":300})

memory = ConversationSummaryBufferMemory(llm= llama_13b, max_token_limit=500)
conversation = ConversationChain(llm= llama_13b, memory=memory,verbose=False)


@st.cache_data(show_spinner=False)
def llama_llm(_llm,prompt):
    response = _llm.predict(prompt)
    return response



@st.cache_data
def process_text(text):
    # Add your custom text processing logic here
    processed_text = text
    return processed_text

@st.cache_resource
def embed(model_name):
    hf_embeddings = HuggingFaceEmbeddings(model_name=model_name)
    return hf_embeddings

@st.cache_data
def chunk_extract(pdf_files):
    pdf_only =[]
    text = ""
    for file in pdf_files:
      if file.endswith('.pdf'):
        pdf_only.append(file)       
      
    merged_pdf = merge_pdfs(pdf_only)
    final_pdf = PyPDF2.PdfReader(merged_pdf)
    for page in final_pdf.pages:
        text += page.extract_text()
      
    for file in pdf_files:
      if file.endswith('xlsx'):
        df = pd.read_excel(file, engine='openpyxl')
        # Find the row index where the table data starts
        data_start_row = 0  # Initialize to 0
        for i, row in df.iterrows():
            if row.notna().all():
                data_start_row = i
                break
              
        if data_start_row>0:  
            df.columns = df.iloc[data_start_row]
          
        
        # Extract the text content above the data
        text += "\n".join(df.iloc[:data_start_row].apply(lambda x: "\t".join(map(str, x)), axis=1)).replace('nan','')
        
        df1 = df.iloc[data_start_row+1:]
        text_buffer = StringIO()
        df1.to_csv(text_buffer, sep='\t', index=False)
        text += "\n\n"+ text_buffer.getvalue()
        text_buffer.close()
        
    chunks_splitted =  text_splitter.split_text(text)
  
    return chunks_splitted


# @st.cache_data
# def embedding_store(pdf_files):
#     pdf_only =[]
#     text = ""
#     for file in pdf_files:
#       if file.endswith('.pdf'):
#         pdf_only.append(file)       
      
#     merged_pdf = merge_pdfs(pdf_only)
#     final_pdf = PyPDF2.PdfReader(merged_pdf)
#     for page in final_pdf.pages:
#         text += page.extract_text()
      
#     for file in pdf_files:
#       if file.endswith('xlsx'):
#         df = pd.read_excel(file, engine='openpyxl')
#         # Find the row index where the table data starts
#         data_start_row = 0  # Initialize to 0
#         for i, row in df.iterrows():
#             if row.notna().all():
#                 data_start_row = i
#                 break
              
#         if data_start_row>0:  
#             df.columns = df.iloc[data_start_row]
          
        
#         # Extract the text content above the data
#         text += "\n".join(df.iloc[:data_start_row].apply(lambda x: "\t".join(map(str, x)), axis=1)).replace('nan','')
        
#         df1 = df.iloc[data_start_row+1:]
#         text_buffer = StringIO()
#         df1.to_csv(text_buffer, sep='\t', index=False)
#         text += "\n\n"+ text_buffer.getvalue()
#         text_buffer.close()
        
#     texts =  text_splitter.split_text(text)
#     docs = text_to_docs(texts)
#     docsearch = FAISS.from_documents(docs, hf_embeddings)
#     return docs, docsearch
    
@st.cache_data
def merge_and_extract_text(pdf_list):
    """
    Helper function to merge PDFs and extract text
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        with open(pdf, 'rb') as file:
            pdf_merger.append(file)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    
    # Extract text from merged PDF
    merged_pdf = PyPDF2.PdfReader(output_pdf)
    all_text = []
    for page in merged_pdf.pages:
        text = page.extract_text()
        all_text.append(text)
    
    return ' '.join(all_text)

def reset_session_state():
    session_state = st.session_state
    session_state.clear()

def get_response(messages: str, model: str = "gpt-3.5-turbo") -> str:
    return openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0.01,
        top_p=0.1,
        #top_k=10,
        seed=1000,
        presence_penalty=0

    )

def wrap_prompt(message: str, role: str) -> dict:
    return {"role": role, "content": message}

def context_data(document):
    prompt_to_add = "Your goal is to identify potential money laundering data from the input transactions data provided by the customer. Output all the potential information that could be related to money laundering or unusual activity, considering past transactions. Strictly output information from the given data only. Do not provide any extra Explanation or Note etc."
    modified_conditions = ['"""' + prompt_to_add + doc + '"""' for doc in document]
    results_textdata = []
    for condition in modified_conditions:
        system_prompt = wrap_prompt("You are a Money Laundering Analyst.", "system")
        user_prompt = wrap_prompt(condition, "user")
        response = get_response([system_prompt, user_prompt])
        results_textdata.append(response['choices'][0]['message']['content'])
    return results_textdata

def context_data_use_llm(document):
    prompt_to_add = "Your goal is to identify potential money laundering data from the input transactions data provided by the customer. Output all the potential information that could be related to money laundering or unusual activity, considering past transactions. Strictly output information from the given data only. Do not provide any extra Explanation or Note etc."
    modified_conditions = ['"""' + prompt_to_add + doc + '"""' for doc in document]
    results_textdata = []
    for condition in modified_conditions:
        
        response = usellm(condition)
        results_textdata.append(response)
    return results_textdata


def calculate_iqr(group):
   
   q=group.quantile(0.75)
   return q

def process_data_credit_card(data_path):
    data=pd.read_excel(data_path,engine='openpyxl')
    
    d = data.dropna(thresh=4)
    d.reset_index(drop=True, inplace=True)
    d.columns = d.iloc[0]
    result_df = d.iloc[1:]
    result_df.columns = [col.replace(" ", "") for col in result_df.columns]
    result1 = result_df[result_df['Debited($)'] >= 5000]

    item_counts = result_df['Description'].value_counts()

    result2 = result_df[result_df['Description'].isin(item_counts.index[item_counts == 1]) & (result_df['Debited($)'] > 1000)]

   

    iqr_val = result_df.groupby('Description')['Debited($)'].apply(calculate_iqr)

    iqr_val = pd.DataFrame(iqr_val)
    iqr_val.reset_index(inplace=True)
    iqr_val.rename(columns={'Debited($)': "IQR"}, inplace=True)

    result = result_df[result_df['Description'].isin(item_counts.index[item_counts > 1])]
    result_df = result.merge(iqr_val, on='Description', how='left')

    result3 = result_df[(result_df['Debited($)'] > result_df["IQR"]) & (result_df['Debited($)'] > 1000)]
    result3=result3.drop(['IQR'],axis=1)
    myDataFrame=pd.concat([result1, result2, result3], axis=0)
    # myDataFrame['Date'] = pd.to_datetime(myDataFrame['Date'])
    # myDataFrame_sorted = myDataFrame.sort_values(by='Date')
    # myDataFrame_sorted.reset_index(inplace=True,drop=True)
    # myDataFrame_sorted['Date'] = myDataFrame_sorted['Date'].dt.strftime('%b %d, %Y')
    json1 = myDataFrame.to_json(orient='records')
    
    return json1



def process_data_saving(data_path):
    data=pd.read_excel(data_path,engine='openpyxl')
    d = data.dropna(thresh=4)
    d.reset_index(drop=True, inplace=True)
    d.columns = d.iloc[0]
    result_df = d.iloc[1:]
    result_df.columns = [col.replace(" ", "") for col in result_df.columns]
    result1 = result_df[result_df['Credited($)'] >= 5000]
    item_counts = result_df['Description'].value_counts()

    result2 = result_df[result_df['Description'].isin(item_counts.index[item_counts == 1]) & (result_df['Credited($)'] > 1000)]

    iqr_val = result_df.groupby('Description')['Credited($)'].apply(calculate_iqr)

    iqr_val = pd.DataFrame(iqr_val)
    iqr_val.reset_index(inplace=True)
    iqr_val.rename(columns={'Credited($)': "IQR"}, inplace=True)

    result = result_df[result_df['Description'].isin(item_counts.index[item_counts > 1])]
    result_df = result.merge(iqr_val, on='Description', how='left')
    result3 = result_df[(result_df['Credited($)'] > result_df["IQR"]) & (result_df['Credited($)'] > 1000)]
    result3=result3.drop(['IQR'],axis=1)
    myDataFrame=pd.concat([result1, result2, result3], axis=0)
    # myDataFrame['Date'] = pd.to_datetime(myDataFrame['Date'])
    # myDataFrame_sorted = myDataFrame.sort_values(by='Date')
    # myDataFrame_sorted.reset_index(inplace=True,drop=True)
    # myDataFrame_sorted['Date'] = myDataFrame_sorted['Date'].dt.strftime('%b %d, %Y')
    json2 = myDataFrame.to_json(orient='records')

    return json2

def append_text_to_file(file_path, text_content):
    try:
        with open(file_path, 'a') as file:
            file.write(text_content)
        print(f"Text has been appended to '{file_path}'.")
    except Exception as e:
        print(f"An error occurred: {e}")

def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        all_text = []
        for page in pdf.pages:
            text = page.extract_text()
            all_text.append(text)
    return "\n".join(all_text)


 


@st.cache_data(show_spinner=False)
def embedding_store_aml(_doc,_hf_embeddings):
    docsearch = FAISS.from_documents(_doc, _hf_embeddings)
    return _doc, docsearch

def embedding_store_aml_2(_doc,_hf_embeddings):
    docsearch = FAISS.from_documents(_doc, _hf_embeddings)
    return _doc, docsearch


# def merge_and_extract_text(pdf_list):
#     merged_pdf = fitz.open()
#     # Merge the PDF files
#     for pdf_file in pdf_list:
#         pdf_document = fitz.open(pdf_file)
#         merged_pdf.insert_pdf(pdf_document)
#     # Create an empty string to store the extracted text
#     merged_text = ""
#     # Extract text from each page of the merged PDF
#     for page_num in range(merged_pdf.page_count):
#         page = merged_pdf[page_num]
#         text = page.get_text()
#         merged_text += text
#     # Close the merged PDF
#     merged_pdf.close()
#     return merged_text


@st.cache_data
def render_pdf_as_images(pdf_file):
    """
    Helper function to render PDF pages as images
    """
    pdf_images = []
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        img = page.get_pixmap()
        img_bytes = img.tobytes()
        pdf_images.append(img_bytes)
    pdf_document.close()
    return pdf_images

# To check if pdf is searchable
def is_searchable_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                return True

    return False



def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        all_text = []
        for page in pdf.pages:
            text = page.extract_text()
            all_text.append(text)
    return "\n".join(all_text)

def replace_strings(input_string, item1, replacement1, item2, replacement2):
    result = input_string.replace(item1, replacement1).replace(item2, replacement2)
    return result






# Function to add checkboxes to the DataFrame
@st.cache_data
def add_checkboxes_to_dataframe(df):
    # Create a new column 'Select' with checkboxes
    checkbox_values = [True] * (len(df) - 1) + [False]  # All True except the last row
    df['Select'] = checkbox_values
    return df

# convert scanned pdf to searchable pdf
def convert_scanned_pdf_to_searchable_pdf(input_file):
    """
     Convert a Scanned PDF to Searchable PDF

    """
    # Convert PDF to images
    print("Running OCR")
    images = convert_from_path(input_file)

    # Preprocess images using OpenCV
    for i, image in enumerate(images):
        # Convert image to grayscale
        image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)

        # Apply thresholding to remove noise
        _, image = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Enhance contrast
        image = cv2.equalizeHist(image)

        # Save preprocessed image
        cv2.imwrite(f'{i}.png', image)

    # Perform OCR on preprocessed images using Tesseract
    text = ''
    for i in range(len(images)):
        image = cv2.imread(f'{i}.png')
        text += pytesseract.image_to_string(image)
    
    return text
## context data extraction:
def context_data(document):
    prompt_to_add = "Your goal is to identify potential money laundering data from the input transactions data provided by the customer. Output only the data that you find related to any money laundering activity.## Strictly output information from the given data only. Do not provide any extra Explanation or Note etc."
    modified_conditions = ['"""' + prompt_to_add + doc + '"""' for doc in document]
    results_textdata = []
    for condition in modified_conditions:
        response = usellm(condition)
        results_textdata.append(response)
    return results_textdata

def process_documents(documents):
    docs_new = []
    for i in documents:
        new_string = replace_strings(i, "\n", "", "\t", "")
        docs_new.append(new_string)
    return docs_new





# Setting globals
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = True
if "stored_session" not in st.session_state:
    st.session_state["stored_session"] = []

if "tmp_table_gpt" not in st.session_state:
    st.session_state.tmp_table_gpt=pd.DataFrame()
if "tmp_table_llama" not in st.session_state:
    st.session_state.tmp_table_llama=pd.DataFrame()
if "tmp_summary_gpt" not in st.session_state:
    st.session_state["tmp_summary_gpt"] = ''
if "tmp_summary_llama" not in st.session_state:
    st.session_state["tmp_summary_llama"] = ''
if "sara_recommendation_gpt" not in st.session_state:
    st.session_state["sara_recommendation_gpt"] = ''

if "tmp_table_gpt_fd" not in st.session_state:
    st.session_state.tmp_table_gpt_fd = pd.DataFrame()
if "tmp_table_llama_fd" not in st.session_state:
    st.session_state.tmp_table_llama_fd = pd.DataFrame()
if "tmp_table_gpt_aml" not in st.session_state:
    st.session_state.tmp_table_gpt_aml = pd.DataFrame()
if "tmp_table_llama_aml" not in st.session_state:
    st.session_state.tmp_table_llama_aml = pd.DataFrame()

if "tmp_summary_gpt_fd" not in st.session_state:
    st.session_state["tmp_summary_gpt_fd"] = ''
if "tmp_summary_llama_fd" not in st.session_state:
    st.session_state["tmp_summary_llama_fd"] = ''
if "tmp_summary_gpt_aml" not in st.session_state:
    st.session_state["tmp_summary_gpt_aml"] = ''
if "tmp_summary_llama_aml" not in st.session_state:
    st.session_state["tmp_summary_llama_aml"] = ''
if "sara_recommendation_gpt_fd" not in st.session_state:
    st.session_state["sara_recommendation_gpt_fd"] = ''
if "sara_recommendation_llama_fd" not in st.session_state:
    st.session_state["sara_recommendation_llama_fd"] = ''

if "sara_recommendation_gpt_aml" not in st.session_state:
    st.session_state["sara_recommendation_gpt_aml"] = ''
if "sara_recommendation_llama_aml" not in st.session_state:
    st.session_state["sara_recommendation_llama_aml"] = ''

if "tmp_narrative_gpt" not in st.session_state:
    st.session_state["tmp_narrative_gpt"] = ''
if "tmp_narrative_llama" not in st.session_state:
    st.session_state["tmp_narrative_llama"] = ''

if "case_num" not in st.session_state:
    st.session_state.case_num = ''
if "fin_opt" not in st.session_state:
    st.session_state.fin_opt = ''
if "context_1" not in st.session_state:
    st.session_state.context_1 = ''
if "llm" not in st.session_state:
    st.session_state.llm = 'Closed-Source'
if "pdf_files" not in st.session_state:
    st.session_state.pdf_files = []

if "lineage_aml" not in st.session_state:
    st.session_state["lineage_aml"] = {}
if "lineage_aml_llama" not in st.session_state:
    st.session_state["lineage_aml_llama"] = {}

if "lineage_gpt" not in st.session_state:
    st.session_state["lineage_gpt"] = {}
if "lineage_gpt_llama" not in st.session_state:
    st.session_state["lineage_gpt_llama"] = {}



# Apply CSS styling to resize the buttons
st.markdown("""
    <style>
        .stButton button {
            width: 145px;
            height: 35px;
        }
    </style>
""", unsafe_allow_html=True)

@st.cache_data
def add_footer_with_fixed_text(doc, footer_text):
    # Create a footer object
    footer = doc.sections[0].footer

    # Add a paragraph to the footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # Set the fixed text in the footer
    paragraph.text = footer_text

    # Add a page number field to the footer
    run = paragraph.add_run()
    fld_xml = f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>'
    fld_simple = parse_xml(fld_xml)
    run._r.append(fld_simple)

    # Set the alignment of the footer text
    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

@st.cache_data
def create_filled_box_with_text(color, text):
    box_html = f'<div style="flex: 1; height: 100px; background-color: {color}; display: flex; align-items: center; justify-content: center;">{text}</div>'
    st.markdown(box_html, unsafe_allow_html=True)

@st.cache_data
def create_zip_file(file_paths, zip_file_name):
    with zipfile.ZipFile(zip_file_name, 'w') as zipf:
        for file_path in file_paths:
            zipf.write(file_path, os.path.basename(file_path))



####### This markdown is to manage app style (app markdown)
st.markdown("""

<style>

.st-d5 {
    line-height: 1;
}


.css-1upf7v9 { 
    gap: 0.5rem;
}

.css-1balh2r{
    gap: 0;
}

.css-1544g2n {
    padding: 0;
    padding-top: 2rem;
    padding-right: 1rem;
    padding-bottom: 1.5rem;
    padding-left: 1rem;
}

.css-1q2g7hi {
    top: 2px;
    min-width: 350px;
    max-width: 600px;
    }

.st-ah {
    line-height: 1;
}

.st-af {
    font-size: 1.5rem;
}

.css-1a65djw {
    gap: 0;
    }

.css-1y4p8pa {
    width: 100%;
    padding-top: 2rem;
    padding-bottom: 10rem;
    max-width: 70rem;
}

.css-xujc5b p{
font-size: 25px;
}


</style>
""", unsafe_allow_html=True)



# Addding markdown styles(Global)
st.markdown("""
<style>
.big-font {
    font-size:60px !important;
}
</style>
""", unsafe_allow_html=True)


# Set Sidebar
st.markdown("""
<style>
    [data-testid=stSidebar] {
        background-color: FBFBFB;
    }
</style>
""", unsafe_allow_html=True)

#Adding llm type-> st.session_state.llm
st.session_state.llm = st.radio("",options = pd.Series(["","Closed-Source","Open-Source"]), horizontal=True)
st.markdown(
    """ <style>
            div[role="radiogroup"] >  :first-child{
                display: none !important;
            }
        </style>
        """,
    unsafe_allow_html=True
                    )
st.markdown("---")

st.title("Suspicious Activity Reporting Assistant")
with st.sidebar:
    # st.sidebar.write("This is :blue[test]")
    # Navbar
    st.markdown('<link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css" integrity="sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm" crossorigin="anonymous">', unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-top navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <style>
    .navbar-brand img {
      max-height: 50px; /* Adjust the height of the logo */
      width: auto; /* Allow the width to adjust based on the height */
    }
    </style>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
            <a class="navbar-brand" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
                <strong>| Operations Process Automation</strong>
            </a>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-bottom navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
        <!--p style='color: white;'><b>Powered by EXL</b></p--!>
        <p style='color: white;'> <strong>Powered by EXL</strong> </p>
            <!--a class="nav-link disabled" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
            </a--!>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    # Add the app name
    st.sidebar.markdown('<p class="big-font">SARA</p>', unsafe_allow_html=True)
    # st.sidebar.header("SARA")
    st.markdown("---")

    # Add a drop-down for case type
    options1 = ["Select Case Type", "Fraud transaction dispute", "Money Laundering"]
    selected_option_case_type = st.sidebar.selectbox("", options1)
    st.markdown("---")
    
    # Add a single dropdown
    options2 = ["Select Case ID", "SAR-2023-24680", "SAR-2023-13579", "SAR-2023-97531", "SAR-2023-86420", "SAR-2023-24681"]
    selected_option = st.sidebar.selectbox("", options2)
    # Add the image to the sidebar below options
    st.sidebar.image("MicrosoftTeams-image (3).png", use_column_width=True)

    
# Assing action to the main section
if selected_option_case_type == "Select Case Type":
    st.header("")

## Fraud Transaction Code started
elif selected_option_case_type == "Fraud transaction dispute":
    st.markdown("### :blue[Fraud transaction dispute]")

# st.markdown('---')

    # Redirect to Merge PDFs page when "Merge PDFs" is selected
    if selected_option == "SAR-2023-24680":
        st.session_state.case_num = "SAR-2023-24680"
        # st.header("Merge Documents")
        # st.write("Upload multiple document files and merge them into one doc.")

        # Upload PDF files
        # st.subheader("Upload Case Files")
        # st.markdown(f"**Case No: {st.session_state.case_num}**")
        # st.markdown("""
        #     | Case No.                  | Case Type                 | Customer Name             | Case Status             | Open Date              |
        #     | ------------------------  | ------------------------- | ------------------------- | ------------------------|------------------------|
        #     | SAR-2023-24680            | Fraud Transaction Dispute | John Brown                | In Progress             | 12/10/2020             |
        #     """)

        col1,col2 = st.columns(2)
        # Row 1
        with col1:
            st.markdown("##### **Case number&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** SAR-2023-24680")
            st.markdown("##### **Customer name  :** John Brown")


        with col2:
            st.markdown("##### **Case open date&nbsp;&nbsp;&nbsp;&nbsp;:** Feb 02, 2021")
            st.markdown("##### **Case type&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Fraud transaction")


        # Row 2
        with col1:
            st.markdown("##### **Customer ID&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** 9659754")


        with col2:
            st.markdown("##### **Case Status&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Open")

        st.markdown("---")

        ## Defining some global variables for fraud transaction

        directoty_path = "data/"
        fetched_files = read_pdf_files(directoty_path)
        


    
        if selected_option:
            
            col1_up, col2_up, col3_up, col4_up, col5_up, col6_up = st.tabs(["Data", "Generate Insights","Lineage","Summarization","Download Report", "Make a Decision"])

            with col1_up:
        
                bt1_up, bt2_up = st.tabs(["Fetch Evidence", "Upload Evidence"])

                with bt1_up:
                    # Set the color
                    # st.markdown(
                    #     """
                    #     <div style="display: flex; justify-content: center; align-items: center; height: 48px; border: 1px solid #ccc; border-radius: 5px; background-color: #f2f2f2;">
                    #         <span style="font-size: 16px;  ">Fetch Evidence</span>
                    #     </div>
                    #     """,
                    #     unsafe_allow_html=True
                    # )
                    if 'clicked' not in st.session_state:
                        st.session_state.clicked = False
                    
                    def set_clicked():
                        st.session_state.clicked = True
                        st.session_state.disabled = True
                    st.write("") #for the gap
                    st.button('Fetch Evidence', on_click=set_clicked)

                    if st.session_state.clicked:
                        # st.write("Evidence Files:") 
                        # st.markdown(html_str, unsafe_allow_html=True)
                        
                        # Showing files
                        # show_files = fetched_files.copy()
                        # show_files = show_files + ['Other.pdf']
                        # files_frame = pd.DataFrame(show_files, columns=["File Name"])
                        # # files_frame["Select"] = [True for _ in range(len(files_frame))]
                        # files_frame = files_frame.reset_index(drop=True)

                        # # Add checkboxes to the DataFrame
                        # df_with_checkboxes = add_checkboxes_to_dataframe(files_frame)
                        
                        # # Iterate through each row and add checkboxes
                        # for index, row in df_with_checkboxes.iterrows():
                        #     if index < len(df_with_checkboxes) - 1:
                        #         checkbox_state = st.checkbox(f" {row['File Name']}", value=True)
                        #         df_with_checkboxes.loc[index, 'Select'] = checkbox_state
                        #     else:
                        #         st.checkbox(f"{row['File Name']}", value=False)



                        # st.dataframe(files_frame)
                        # st.write(df_reset.to_html(index=False), unsafe_allow_html=True)
                        # st.markdown(files_frame.style.hide(axis="index").to_html(), unsafe_allow_html=True)
                        
                        
                        
                        #select box to select file
                        selected_file_name = st.selectbox(":blue[Select a file to View]",fetched_files)
                        st.write("Selected File: ", selected_file_name)
                        st.session_state.disabled = False
                        file_ext = tuple("pdf")
                        if selected_file_name.endswith(file_ext):
                            selected_file_path = os.path.join(directoty_path, selected_file_name)
                            #converting pdf data to bytes so that render_pdf_as_images could read it
                            file = pdf_to_bytes(selected_file_path)
                            pdf_images = render_pdf_as_images(file)
                            #showing content of the pdf
                            st.subheader(f"Contents of {selected_file_name}")
                            for img_bytes in pdf_images:
                                st.image(img_bytes, use_column_width=True)
                        else:
                            selected_file_path = os.path.join(directoty_path, selected_file_name)
                            # This is showing png,jpeg files
                            st.image(selected_file_path, use_column_width=True)



                with bt2_up:
                    pdf_files = st.file_uploader("", type=["pdf","png","jpeg","docx","xlsx"], accept_multiple_files=True)
                    st.session_state.pdf_files = pdf_files
                    # showing files
                    for uploaded_file in pdf_files:
                        #This code is to show pdf files
                        file_ext = tuple("pdf")
                        if uploaded_file.name.endswith(file_ext):
                            # Show uploaded files in a dropdown
                            # if pdf_files:
                            st.subheader("Uploaded Files")
                            file_names = [file.name for file in pdf_files]
                            selected_file = st.selectbox(":blue[Select a file]", file_names)
                            # Enabling the button
                            st.session_state.disabled = False
                            # Display selected PDF contents
                            if selected_file:
                                selected_pdf = [pdf for pdf in pdf_files if pdf.name == selected_file][0]
                                pdf_images = render_pdf_as_images(selected_pdf)
                                st.subheader(f"Contents of {selected_file}")
                                for img_bytes in pdf_images:
                                    st.image(img_bytes, use_column_width=True)

                        else:
                            # This is showing png,jpeg files
                            st.image(uploaded_file, use_column_width=True)

            #creating temp directory to have all the files at one place for accessing
                # tmp_dir_ = tempfile.mkdtemp()
                # temp_file_path= []


                # for uploaded_file in pdf_files:
                #     file_ext = tuple("pdf")
                #     if uploaded_file.name.endswith(file_ext):
                #         file_pth = os.path.join(tmp_dir_, uploaded_file.name)
                #         with open(file_pth, "wb") as file_opn:
                #             file_opn.write(uploaded_file.getbuffer())
                #             temp_file_path.append(file_pth)
                #     else:
                #         pass

                # for fetched_pdf in fetched_files:
                #     file_ext = tuple("pdf")
                #     if fetched_pdf.endswith(file_ext):
                #         file_pth = os.path.join('data/', fetched_pdf)
                #         # st.write(file_pth)
                #         temp_file_path.append(file_pth) 
                #     else:
                #         pass   


            temp_file_path = pytesseract_code(directoty_path,fetched_files)        

            with col2_up:
                #This is the embedding model
                model_name = "sentence-transformers/all-MiniLM-L6-v2"
                # model_name = "hkunlp/instructor-large"
                
                # Memory setup for gpt-3.5
                llm = ChatOpenAI(temperature=0.1)
                memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=500)
                conversation = ConversationChain(llm=llm, memory =memory,verbose=False)
                
                
                # Adding condition on embedding
                try:
                    if temp_file_path:
                        hf_embeddings = embed(model_name) 
                    else:
                        pass
                except NameError:
                    pass
                
                # Chunking with overlap
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size = 1000,
                    chunk_overlap  = 100,
                    length_function = len,
                    separators=["\n\n", "\n", " ", ""]
                )
               

               # Creating header
                col1,col2 = st.columns(2)
                with col1:
                    st.markdown("""<span style="font-size: 24px; ">Key-Questions</span>""", unsafe_allow_html=True)
                    # Create a Pandas DataFrame with your data
                    data = {'Questions': ["What is the customer's name?","What is the suspect's name?",' List the merchant name',' How was the bank notified?',' When was the bank notified?',' What is the fraud type?',' When did the fraud occur?',' Was the disputed amount greater than 5000 USD?','what type of network/card is used in transaction?',' Was the police report filed?']}
                    df_fixed = pd.DataFrame(data)
                    df_fixed.index = df_fixed.index +1
                with col2:
                    # Create a checkbox to show/hide the table
                    cols1, cols2, cols3, cols4 = st.columns([1,1,1,1])
                    with cols1:
                        show_table1 = tog.st_toggle_switch(label="", 
                                            key="Key1", 
                                            default_value=False, 
                                            label_after = False, 
                                            inactive_color = '#D3D3D3', 
                                            active_color="#11567f", 
                                            track_color="#29B5E8"
                                            )
                    # Show the table if the checkbox is ticked
                    if show_table1:
                        df_fixed["S.No."] = df_fixed.index
                        df_fixed = df_fixed.loc[:,['S.No.','Questions']]
                        st.markdown(df_fixed.style.hide(axis="index").to_html(), unsafe_allow_html=True)

                with st.spinner('Wait for it...'):
                    if 'clicked1' not in st.session_state:
                    
                       st.session_state.clicked1 = False
                
                    def set_clicked1():
                        st.session_state.clicked1 = True
                        st.session_state.disabled = True
                    st.button("Generate Insights",key=2,on_click=set_clicked1,disabled=st.session_state.disabled)
                    if st.session_state.clicked1:
                        if temp_file_path is not None:
                        # File handling logic
                            _, docsearch = embedding_store(temp_file_path,hf_embeddings)
                            res_dict = {}
                            lineage_dict = {}
                
                            if st.session_state.llm == "Closed-Source":
                                query = "What is the customer's name?"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f'''Perform Name Enitity Recognition to identify the cardholder name as accurately as possible, given the context. The customer can also be referenced as the cardholder with whom the Fraud has taken place.\n\n\ 
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''
                                response = usellm(prompt_1)
                                res_dict[query] = response
                                lineage_dict[query] = context_1
                                st.session_state["lineage_gpt"][query] = context_1
                            
                                    
                                query = "What is the suspect's name?"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f'''Perform Name Enitity Recognition to identify the suspect name as accurately as possible, given the context. Suspect is the Person who has committed the fraud with the Customer. Respond saying "The Suspect Name is not Present" if there is no suspect in the given context.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''
                                response = usellm(prompt_1)
                                res_dict[query] = response
                                lineage_dict[query] = context_1
                                st.session_state["lineage_gpt"][query] = context_1
                        

                                        
                                query = "List the Merchant Name"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f'''Perform Name Enitity Recognition to identify all the Merchant Organizations as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''
                                response = usellm(prompt_1)
                                res_dict[query] = response
                                lineage_dict[query] = context_1
                                st.session_state["lineage_gpt"][query] = context_1


                                    
                                query = "How was the bank notified?"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f''' You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Give a concise response in one sentence.)'''
                                response = usellm(prompt_1)
                                res_dict[query] = response
                                lineage_dict[query] = context_1
                                st.session_state["lineage_gpt"][query] = context_1
                            


                                query = "When was the bank notified?"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f''' You need to act as a Financial analyst to Identify the processing date when bank was notified of the fraud?. Given the context, provide a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''
                                response = usellm(prompt_1)
                                res_dict[query] = response  
                                lineage_dict[query] = context_1  
                                st.session_state["lineage_gpt"][query] = context_1           

                                    
                                query = "What type of fraud is taking place?"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Prvide me a concise and relevant response)'''
                                response = usellm(prompt_1)
                                res_dict[query] = response  
                                lineage_dict[query] = context_1
                                st.session_state["lineage_gpt"][query] = context_1
                                
                                query = "When did the fraud occur?"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f''' You need to act as a Financial analyst to identify the when the did the fraud occur i.e., the Dispute Date. Given the context, provide a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''
                                response = usellm(prompt_1)
                                res_dict[query] = response 
                                lineage_dict[query] = context_1 
                                st.session_state["lineage_gpt"][query] = context_1
                                        
                                query = "Was the disputed amount greater than 5000 usd?"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 or no, given the context. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''
                                response = usellm(prompt_1)
                                res_dict[query] = response  
                                lineage_dict[query] = context_1
                                st.session_state["lineage_gpt"][query] = context_1

                                    
                                query = "What type of network/card is used in transaction?"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f''' You need to act as a Financial analyst to identify the type of card and card's brand involved, given the context. On a higher level the card can be a Credit or Debit Card. VISA, MasterCard or American Express, Citi Group, etc. are the different brands with respect to a Credit card or Debit Card . Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Prvide me a concise and relevant response)'''
                                response = usellm(prompt_1)
                                res_dict[query] = response  
                                lineage_dict[query] = context_1
                                st.session_state["lineage_gpt"][query] = context_1
                            

                                            
                                query = "Was the police report filed?"
                                context_1 = docsearch.similarity_search(query, k=5)
                                prompt_1 = f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Prvide me a concise and relevant response)'''
                                response = usellm(prompt_1)
                                res_dict[query] = response 
                                lineage_dict[query] = context_1
                                st.session_state["lineage_gpt"][query] = context_1

                                #st.session_state["lineage_gpt"] = lineage_dict


                
                                try:
                                    #resp_dict_obj = json.loads(response)
                                    res_df_gpt = pd.DataFrame(res_dict.items(), columns=['Question','Answer'])
                                    res_df_gpt.reset_index(drop=True, inplace=True)
                                    index_ = pd.Series([1,2,3,4,5,6,7,8,9,10])
                                    res_df_gpt = res_df_gpt.set_index([index_])   
                                except:
                                    e = Exception("")
                                    st.exception(e)

                               

                                st.table(res_df_gpt)

                                # tmp_table_gpt = res_df_gpt

                                st.session_state["tmp_table_gpt"] = pd.concat([st.session_state.tmp_table_gpt, res_df_gpt], ignore_index=True)
                                
                                query ="Is invoice is billed to cardholder or someone else?"
                                contexts = docsearch.similarity_search(query, k=9)
                                prompt = f" You are professional Fraud Analyst. Find answer to the questions as truthfully and in as detailed as possible as per given context only,\n\n\
                                cardholder's name,adress can be identified from cardholder information. Cardholder is the person who is the owner of the card, cardholder can also be referenced as the victim with whom fraud has taken place.\n\n\
                                Identify to whom invoice is billed (Detials mentioned in invoice is of the person who made the transaction,it may be or may not be of the cardholder)\n\n\
                                Compare both the details, if details mentioned in invoice matches the cardholder details, then invoice is billed to cardholder else it is billed to someone else who misued the card.\n\n\
                                    Context: {contexts}\n\
                                    Response (Give me a concise response.)"
                                response_1 = usellm(prompt)



                                query ="Give your recommendation if this is a Suspicious activity or not?"
                                contexts = docsearch.similarity_search(query, k=9)
                                prompt = f" You are professional Fraud Analyst. Find answer to the questions as truthfully and in as detailed as possible as per given context only,\n\n\
                                    1. Check if The transaction/disputed amount > 5,000 USD value threshold, If Yes, then check below points to make sure if it is a suspicious activity or not: \n\
                                    2. {response_1} analyse this response,if details matches or not? If matches then there is no suspicion else, it can be a suspicipos activity. (Also mention the mismatched details).\n\n\
                                    3. If a potential suspect name is identified or not? Suspect is a person who has commited the fraud, If identified then this can be a suspicious activity, else not.\n\n\
                                    Even if transaction/disputed amount > 5,000 USD but if above criteria does not met, then this can not be considered as a suspicious activity. \n\n\
                                    Based on above points, give your recommendation if this is a case of suspicious activity or not? \n\n\
                                    Context: {contexts}\n\
                                    Response (Give me a concise recommendation in few pointers.If not a case of suspicion them mention it properly.)"
                                response1 = usellm(prompt) 
                                
                                # This replace text is basically to stop rendering of $ to katex (that creates the text messy, hence replacing $)
                                response1 = response1.replace("$", " ")
                                response1 = response1.replace("5,000", "5,000 USD")
                                response1 = response1.replace("5,600", "5,600 USD")
                                st.session_state["sara_recommendation_gpt"] = response1  
                                # sara_recommendation_gpt = response1 
                                        
                                
                                st.markdown("### SARA Recommendation")
                                st.markdown(response1)

                                
                                st.markdown("#### Recommendation Feedback:")
                                col_1, col_2, col_3, col_4, col_5, col_6 = st.columns(6)

                                with col_1:
                                    if st.button("👍🏻",key=3):
                                        st.write("*Feedback is recorded*")
                        

                                with col_2:
                                    if st.button("👎🏻",key=4):
                                        st.write("*Feedback is recorded*")

                                                        
                            
                                    

                                
                            elif st.session_state.llm == "Open-Source":

                                chat_history = {}

                                query = "What is the customer's name?"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 = f'''You are a professional fraud analyst. Perform Name Enitity Recognition to identify the victim's name as accurately as possible, given the context. The victim can also be referenced as the customer with whom the Fraud has taken place.
                                victim's name is the Name provided in Cardholder Information.\n\n\
                                        Question: {query}\n\
                                        Context: {context_1}\n\
                                        Response: (Give me response in one sentence. Do not give me any Explanation or Note)'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response


                                query = "What is the suspect's name?"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 =  f'''Act as a professional fraud analyst.You need to check the document and compare if any name discrepencies are present that points towards the suspect who used the card without the consent of the cardholder.
                                            Take the provided information as accurate. Reply the name of the person who is the suspect. \n\n\
                                            Context: {context_1}\n\
                                            Response: (Give a short response in a single sentence.Do not add any extra Information, Explanation,Note.)'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response

                                
                                
                                query = "List the merchant name"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 = f'''You are a professional fraud analyst, perform Name Enitity Recognition to identify Merchant as accurately as possible from the provided information.A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and short response.\n\n\
                                Take the provided information as accurate.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Give a short response in a single sentence. Do not add any extra Information,Explanation,Note.)'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response


                                query = "How was the bank notified?"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 =  f'''You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Give me a concise response in one sentence. Do not give me any further Explanation, Note )'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response

                                
                                query = "When was the bank notified?"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 =  f'''You need to act as a Financial analyst to identify when the bank was notified of the Fraud. Look for the disputed date. Given the context, provide a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Give me a concise response in one sentence.Do not add any prefix like 'Response' or 'Based on the document'. Do not add any extra Explanation, Note)'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response
                                


                                query = "What is the Fraud Type?"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 =  f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Give me response in one sentence. Do not add prefix like 'Response' or 'based on the document. Do not give me any Explanation or Note)'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response




                                query = "When did the fraud occur?"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 =  f''' You need to act as a Financial analyst to identify the when the did the fraud occur i.e., the Transaction Date. Given the context, provide a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Give me a concise response in one sentence. Do not add prefix like 'based on the document. Do not add any further Explanation or Note.)'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response


                                query = "Was the disputed amount greater than 5000 usd?"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 =  f''' You need to act as a Financial analyst to identify the disputed amount.Perform a mathematical calculation to identify if the disputed amount is greater than 5000 USD or not.Given the context, give a relevant and concise response.\n\n\
                                                Take the provided information as accurate. \n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Give a short response in a single sentence. Do not give any extra Explanation, Note, Descricption, Information.)'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response


                                query = "What type of cards are involved?"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 =  f''' You need to act as a Financial analyst to identify the type of card and card network involved, given the context. On a higher level the card can be a Credit Visa, Debit Visa Card.Based on the context give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Give me a concise response in one sentence.Do not add prefix like: ['based on the document']. Do not add any further Explanation, Note.')'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response


                                query = "Was the police report filed?"
                                context_1 = docsearch.similarity_search(query, k=9)
                                prompt_1 =  f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: (Provide a concise Response in a single sentence. Do not write any extra [Explanation, Note, Descricption].)'''
                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response

                                try:
                                    res_df_llama = pd.DataFrame(list(chat_history.items()), columns=['Question','Answer'])
                                    res_df_llama.reset_index(drop=True, inplace=True)
                                    index_ = pd.Series([1,2,3,4,5,6,7,8,9,10])
                                    res_df_llama = res_df_llama.set_index([index_])
                                    # st.write(res_df_llama)
                                except IndexError: 
                                    pass
                                st.table(res_df_llama)
                                st.session_state["tmp_table_llama"] = pd.concat([st.session_state.tmp_table_llama, res_df_llama], ignore_index=True)


                                st.write("#### SARA Recommendation")
                                
                                st.markdown("""<span style="font-size: 16px;">Based on the given context, the activity can be considered suspicious due to the following reasons:</span>""", unsafe_allow_html=True)
                                st.markdown("""<span style="font-size: 16px;">1. The transaction amount is $5,600.</span>""", unsafe_allow_html=True)
                                st.markdown("""<span style="font-size: 16px;">2. The fraud type is mentioned as "83 Fraud - Card Absent Environment,which indicates the absence of the physical card during the transaction.</span>""", unsafe_allow_html=True)
                                st.markdown("""<span style="font-size: 16px;">3. A suspect named Mike White is identified, address is mentioned as 520 Wintergreen Ct, Vancaville, CA 95587 in the merchant invoice,further confirming the possibility of fraud.</span>""", unsafe_allow_html=True)
                                st.markdown("""<span style="font-size: 16px;">Considering the above findings, the activity can be considered as suspicious and should be investigated further</span>""", unsafe_allow_html=True)

                                if st.session_state.clicked1:

                                    st.markdown("#### Recommendation Feedback:")
                                    col_1, col_2, col_3, col_4, col_5, col_6 = st.columns(6)

                                    with col_1:
                                        if st.button("👍🏻",key=3):
                                            st.write("*Feedback is recorded*")
                                        # st.markdown('<span style="font-size: 24px;">👍🏻</span>',unsafe_allow_html=True)
                            

                                    with col_2:
                                        if st.button("👎🏻",key=4):
                                            st.write("*Feedback is recorded*")
                                        # st.markdown('<span style="font-size: 24px;">👎🏻</span>',unsafe_allow_html=True)
    

                               
                            
                               




                st.markdown("---")

                # For input box outside of template4
                try:
                    if temp_file_path:
                        docs, docsearch = embedding_store(temp_file_path,hf_embeddings)
                    else:
                        pass
                except Exception:
                    pass


                # Text Input
                # st.markdown("""<span style="font-size: 24px; ">Ask Additional Questions</span>""", unsafe_allow_html=True)
                query = st.text_input(':black[Ask Additional Questions]',disabled=st.session_state.disabled)
                text_dict = {}

                @st.cache_data
                def LLM_Response():
                    llm_chain = LLMChain(prompt=prompt, llm=llm)
                    response = llm_chain.run({"query":query, "context":context})
                    return response

                if st.session_state.llm == "Closed-Source":
                    with st.spinner('Getting you information...'):      
                        if query:
                            # Text input handling logic
                            #st.write("Text Input:")
                            #st.write(text_input)

                            context_1 = docsearch.similarity_search(query, k=9)
                            st.session_state.context_1 = context_1
                            if query.lower() == "What is the customer's name?":
                                prompt_1 = f'''Perform Name Enitity Recognition to identify the Customer name as accurately as possible, given the context. The Customer can also be referenced as the Victim or the person with whom the Fraud has taken place.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                
                            elif query.lower() == "What is the suspect's name?":
                                prompt_1 = f'''Perform Name Enitity Recognition to identify the Suspect name as accurately as possible, given the context. Suspect is the Person who has committed the fraud with the Customer. Respond saying "The Suspect Name is not Present" if there is no suspect in the given context.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                
                            elif query.lower() == "List the merchant name":
                                prompt_1 = f'''Perform Name Enitity Recognition to identify all the Merchant Organizations as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                
                            elif query.lower() == "How was the bank notified?":
                                prompt_1 = f''' You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                
                            elif query.lower() == "When was the bank notified?":
                                prompt_1 = f''' You need to act as a Financial analyst to identify the when the bank was notified of the Fraud i.e., the disputed date. Given the context, provide a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                
                            elif query.lower() == "What type of fraud is taking place?":
                                prompt_1 = f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                            
                            elif query.lower() == "When did the fraud occur?":
                                prompt_1 = f''' You need to act as a Financial analyst to identify the when the did the fraud occur i.e., the Transaction Date. Given the context, provide a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                    
                            elif query.lower() == "Was the disputed amount greater than 5000 usd?":
                                prompt_1 = f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 or no, given the context. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                
                            elif query.lower() == "What type of network/card is used in transaction?":
                                prompt_1 = f''' You need to act as a Financial analyst to identify the type of card and card's brand involved, given the context. On a higher level the card can be a Credit or Debit Card. VISA, MasterCard or American Express, Citi Group, etc. are the different brands with respect to a Credit card or Debit Card . Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                
                            elif query.lower() == "Was the police report filed?":
                                prompt_1 = f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                    
                            elif query.lower() == "Is this a valid SAR case?":
                                prompt_1 = f''' You need to act as a Financial analyst to check if this is a SAR or not, given the following context, if the transaction amount is less than 5000 USD we cannot categorize this as SAR (Suspicious activity Report).Give a relevant and concise response. \n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\
                                            Response: '''

                                
                            else:
                                prompt_1 = f'''Act as a financial analyst and give concise answer to below Question as truthfully as possible, with given Context.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\                      
                                            Response: '''

                                        #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                            response = usellm(prompt_1) #LLM_Response()
                            text_dict[query] = response
                            # resp_dict_obj.update(text_dict)
                            st.write(response)
                            if response:
                                df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                            else:
                                df = pd.DataFrame()

                            st.session_state["tmp_table_gpt"] = pd.concat([st.session_state.tmp_table_gpt, df], ignore_index=True)
                            st.session_state.tmp_table_gpt.drop_duplicates(subset=['Question'])
                
                    # #Lineage
                    # retriever(temp_file_path,hf_embeddings)


                elif st.session_state.llm == "Open-Source":
                        
                        with st.spinner('Getting you information...'):      
                            if query:
                                # Text input handling logic
                                #st.write("Text Input:")
                                #st.write(text_input)

                                context_1 = docsearch.similarity_search(query, k=9)
                                st.session_state.context_1 = context_1
                                if query.lower() == "What is the customer's name?":
                                    prompt_1 = f'''Perform Name Enitity Recognition to identify the Customer name as accurately as possible, given the context. The Customer can also be referenced as the Victim or the person with whom the Fraud has taken place.
                                                Customer/Victim is cardholder, whose card is used without their consent.
                                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response: (Provide a concise Response.) '''

                                    
                                elif query.lower() == "What is the suspect's name?":
                                    prompt_1 = f''''Perform Name Enitity Recognition to identify the Suspect name as accurately as possible, given the context. Suspect is the Person who has committed the fraud with the Customer. Respond saying "The Suspect Name is not Present" if there is no suspect in the given context.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response: (Give me response in one sentence.Do not give me any Explanation or Note)'''


                                    
                                elif query.lower() == "List the merchant name":
                                    prompt_1 = f'''Perform Name Enitity Recognition to identify all the Merchant Organizations as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.
                                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''

                                    
                                elif query.lower() == "How was the bank notified?":
                                    prompt_1 = f''' You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a relevant and concise response.
                                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response:(Provide a concise Response.) '''

                                    
                                elif query.lower() == "When was the bank notified?":
                                    prompt_1 = f''' You need to act as a Financial analyst to identify the when the bank was notified of the Fraud i.e., the disputed date. Given the context, provide a relevant and concise response.
                                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response: (Provide a concise Response.)'''

                                    
                                elif query.lower() == "What type of fraud is taking place?":
                                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.
                                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''

                                
                                elif query.lower() == "When did the fraud occur?":
                                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of card and card network involved, given the context. On a higher level the card can be a Credit Visa, Debit Visa Card.Based on the context give a relevant and concise response..
                                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''

                                        
                                elif query.lower() == "Was the disputed amount greater than 5000 usd?":
                                    prompt_1 = f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 or no, given the context. Give a relevant and concise response.
                                                Kindly do not provide any extra [Explanation, Note, Description] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response:(Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.) '''

                                    
                                elif query.lower() == "What type of network/card is used in transaction?":
                                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of Card and Card Network involved, given the context. On a higher level the card can be a Dedit, Crebit Card. VISA, MasterCard, American Express, Citi Group, etc. are the different brands with respect to a Credit Card or Debit Card . Give a relevant and concise response.
                                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response:(Act like a professional and provide me a concise Response . Do not add any extra [Explanation, Note, Descricption] below the context.) '''

                                    
                                elif query.lower() == "Was the police report filed?":
                                    prompt_1 = f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.
                                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''

                                elif query.lower() == "Is this a valid sar case?":
                                    prompt_1 =  f''' You are a Fraud Analyst.Check if there is evidence for this case to address as SAR or not. A SAR case is a case of financial Suspicious/Fraud Activity which can be observed given the context.
                                                If there is any activity without the consent of the cardholder, also if there is a suspect who used the card without the consent.
                                                Then we can address this as a valid SAR case.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response: (Provide a concise response in single sentence.Do not add prefix like ['Respone', 'based on the document']. Do not add any further Explanation,Note.)'''        
                                
                                
                                elif query.lower() == "Is there any evidence of a sar case?":
                                    prompt_1 = f''' You are a Fraud Analyst.Check if there is evidence for this case to address as SAR or not. A SAR case is a case of financial Suspicious/Fraud Activity which can be observed given the context.
                                                If there is any activity without the consent of the cardholder, also if there is a suspect who used the card without the consent.
                                                Then we can address this as a SAR case.Give a concise response with the suspect name. \n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\
                                                Response:(Do not add prefix like ['Respone', 'based on the document']. Do not add any further Explanation,Note.) '''

                                    
                                else:
                                    prompt_1 = f'''Act as a financial analyst and give concise answer to below Question as truthfully as possible, with given Context.
                                                Do not provide any extra [Explanation, Note,Description] block below the Response.\n\n\
                                                Question: {query}\n\
                                                Context: {context_1}\n\                      
                                                Response: (Act like a professional and provide me a concise Response . Do not add any extra [Explanation, Note, Descricption] below the Response.)'''
    

                            #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                                response = llama_llm(llama_13b,prompt_1)
                                text_dict[query] = response

                                st.write(response)

                                if response:
                                    df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                                else:
                                    df = pd.DataFrame()

                                st.session_state["tmp_table_llama"] = pd.concat([st.session_state.tmp_table_llama, df], ignore_index=True)
                                st.session_state.tmp_table_llama.drop_duplicates(subset=['Question'])

            with col3_up:
                if st.session_state["lineage_gpt"] is not None:
 
                    li = ["Select question to get the lineage","What is the customer's name?","What is the suspect's name?","List the Merchant Name","How was the bank notified?","When was the bank notified?","What type of fraud is taking place?","When did the fraud occur?","Was the disputed amount greater than 5000 usd?","What type of network/card is used in transaction?","Was the police report filed?"]
                    
                   
                    selected_option = st.selectbox("", li)
                    if selected_option in li[1:]:
                        doc = st.session_state["lineage_gpt"][selected_option]
                        for i in range(len(doc)):
                            #st.write(doc[i])
                            y=i+1
                            st.write(f":blue[Chunk-{y}:]")
                            st_ = doc[i].page_content.replace("($)"," ")
                            st.write(":blue[Page Content:]",st_) 
                            st.write(":blue[Source:]",doc[i].metadata['source'])
                                   
            with col4_up:
                def summ_gpt_(tmp_table_gpt):
                    template = """Write a concise summary of the context provided.
                    ```{text}```
                    Response: (Return your response in a single paragraph. Please don't include words like these: 'chat summary', 'includes information', 'repetitions of information',' repetitive information'  in my final summary.) """
                    prompt = PromptTemplate(template=template,input_variables=["text"])
                    llm_chain_gpt = LLMChain(prompt=prompt,llm=llm)

                    summ_dict_gpt = tmp_table_gpt.set_index('Question')['Answer']
                    # st.write(summ_dict_gpt)
                    text = []
                    for key,value in summ_dict_gpt.items():
                        text.append(value)
                    response_summ_gpt = llm_chain_gpt.run(text)
                    # st.write(response_summ_gpt)
                    return response_summ_gpt,summ_dict_gpt

                if 'clicked2' not in st.session_state:
                    st.session_state.clicked2 = False
                
                def set_clicked2():
                    st.session_state.clicked2 = True
                    st.session_state.disabled = True
                st.markdown("""<span style="font-size: 24px; ">Summarize key findings of the case.</span>""", unsafe_allow_html=True)
                st.write()
                st.button("Summarize",on_click=set_clicked2,disabled=st.session_state.disabled)    
                with st.spinner("Summarize...."):
                    if st.session_state.clicked2:

                        if st.session_state.llm == "Closed-Source":
                            st.session_state.disabled=False
                            summ_dict_gpt = st.session_state.tmp_table_gpt #.set_index('Question')['Answer'].to_dict()
                            # chat_history = resp_dict_obj['Summary']
                            response_summ_gpt,summ_dict_gpt = summ_gpt_(summ_dict_gpt)
                            response_summ_gpt = response_summ_gpt.replace("$", " ")
                            response_summ_gpt = response_summ_gpt.replace("5,000", "5,000 USD")
                            response_summ_gpt = response_summ_gpt.replace("5,600", "5,600 USD")
                            # memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=300)
                            # memory.save_context({"input": "This is the entire summary"}, 
                            #                 {"output": f"{summ_dict_gpt}"})
                            # conversation = ConversationChain(
                            # llm=llm, 
                            # memory = memory,
                            # verbose=True)
                            st.session_state["tmp_summary_gpt"] = response_summ_gpt
                            # st.session_state["tmp_summary_gpt"] = conversation.predict(input="Provide a detailed summary of the text provided by reframing the sentences. Provide the summary in a single paragraph. Please don't include words like these: 'chat summary', 'includes information' in my final summary.")
                            # st.session_state["tmp_summary_gpt"] = st.session_state["tmp_summary_gpt"].replace("$", "USD ")
                           
                            # showing the text in a textbox
                            # usr_review = st.text_area("", value=st.session_state["tmp_summary_gpt"])
                            # if st.button("Update Summary"):
                            #     st.session_state["fin_opt"] = usr_review
                            st.write(st.session_state["tmp_summary_gpt"])


                        elif st.session_state.llm == "Open-Source":
                            st.session_state.disabled=False
                            template = """Write a detailed summary.
                            Return your response in a single paragraph.
                            ```{text}```
                            Response: """
                            prompt = PromptTemplate(template=template,input_variables=["text"])
                            llm_chain_llama = LLMChain(prompt=prompt,llm=llama_13b)

                            summ_dict_llama = st.session_state.tmp_table_llama.set_index('Question')['Answer']
                            text = []
                            for key,value in summ_dict_llama.items():
                                text.append(value)
                            st.session_state["tmp_summary_llama"] = llm_chain_llama.run(text)
                            st.write(st.session_state["tmp_summary_llama"])

                    if st.session_state.clicked2:
                        st.markdown("#### Summarization Feedback:")
                        col_1, col_2, col_3, col_4, col_5,col_6 = st.columns(6)

                        with col_1:
                            if st.button("👍🏻",key=5):
                                st.write("*Feedback is recorded*")
                            # st.markdown('<span style="font-size: 24px;">👍🏻</span>',unsafe_allow_html=True)


                        with col_2:
                            if st.button("👎🏻",key=6):
                                st.write("*Feedback is recorded*")
                            # st.markdown('<span style="font-size: 24px;">👎🏻</span>',unsafe_allow_html=True)

                    
                tmp_summary = []
                tmp_table = pd.DataFrame()
                try:

                    if st.session_state.llm == "Closed-Source":
                        st.session_state.disabled=False
                        tmp_summary.append(st.session_state["tmp_summary_gpt"])
                        tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_gpt"]], ignore_index=True)
                        tmp_table.drop_duplicates(inplace=True)
                    

                    elif st.session_state.llm == "Open-Source":
                        st.session_state.disabled=False
                        tmp_summary.append(st.session_state["tmp_summary_llama"])
                        tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_llama"]], ignore_index=True)
                        tmp_table.drop_duplicates(inplace=True)

                except:
                    e = Exception("")
                    st.exception(e)


                try:
                    # initiate the doc file
                    doc = docx.Document()
                    # doc.add_section(WD_SECTION.NEW_PAGE)
                    doc.add_heading(f"Case No.: {st.session_state.case_num}",0)

                    # Add a subheader for case details
                    subheader_case = doc.add_paragraph("Case Details")
                    subheader_case.style = "Heading 2"
                    # Addition of case details
                    paragraph = doc.add_paragraph(" ")
                    case_info = {
                        "Case Number                            ": " SAR-2023-24680",
                        "Customer Name                       ": " John Brown",
                        "Customer ID                              ": " 9659754",
                        "Case open date                         ": " Feb 02, 2021",
                        "Case Type                                  ": " Fraud Transaction",
                        "Case Status                                ": " Open"
                    }
                    for key_c, value_c in case_info.items():
                        doc.add_paragraph(f"{key_c}: {value_c}")
                    paragraph = doc.add_paragraph(" ")

                    # Add a subheader for customer info to the document ->>
                    subheader_paragraph = doc.add_paragraph("Customer Information")
                    subheader_paragraph.style = "Heading 2"
                    paragraph = doc.add_paragraph(" ")

                    # Add the customer information
                    customer_info = {
                        "Name                                           ": " John Brown",
                        "Address                                      ": " 858 3rd Ave, Chula Vista, California, 91911 US",
                        "Phone                                          ": " (619) 425-2972",
                        "A/C No.                                        ": " 4587236908230087",
                        "SSN                                               ": " 653-30-9562"
                    }

                    for key, value in customer_info.items():
                        doc.add_paragraph(f"{key}: {value}")
                    paragraph = doc.add_paragraph()
                    # Add a subheader for Suspect infor to the document ->>
                    subheader_paragraph = doc.add_paragraph("Suspect's Info")
                    subheader_paragraph.style = "Heading 2"
                    paragraph = doc.add_paragraph()
                    #""" Addition of a checkbox where unticked box imply unavailability of suspect info"""

                    # Add the customer information
                    sent_val = "Suspect has been reported."
                    paragraph = doc.add_paragraph()
                    runner = paragraph.add_run(sent_val)
                    runner.bold = True
                    runner.italic = True
                    suspect_info = {
                        "Name                                             ": "Mike White",
                        "Address                                        ": "520, WintergreenCt,Vancaville,CA,95587",
                        "Phone                                             ": "NA",
                        "SSN                                                 ": "NA",
                        "Relationship with Customer  ": "NA"
                    }

                    for key, value in suspect_info.items():
                        doc.add_paragraph(f"{key}: {value}")

                    doc.add_heading('Summary', level=2)
                    paragraph = doc.add_paragraph()
                    doc.add_paragraph(tmp_summary)
                    paragraph = doc.add_paragraph()
                    doc.add_heading('Key Insights', level=2)
                    paragraph = doc.add_paragraph()
                    columns = list(tmp_table.columns)
                    table = doc.add_table(rows=1, cols=len(columns), style="Table Grid")
                    table.autofit = True
                    for col in range(len(columns)):
                        # set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50) # set cell margin
                        table.cell(0, col).text = columns[col]
                    # doc.add_table(st.session_state.tmp_table.shape[0]+1, st.session_state.tmp_table.shape[1], style='Table Grid')

                    for i, row in enumerate(tmp_table.itertuples()):
                        table_row = table.add_row().cells # add new row to table
                        for col in range(len(columns)): # iterate over each column in row and add text
                            table_row[col].text = str(row[col+1]) # avoid index by adding col+1
                    # save document
                    # output_bytes = docx.Document.save(doc, 'output.docx')
                    # st.download_button(label='Download Report', data=output_bytes, file_name='evidence.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

                    paragraph = doc.add_paragraph()
                    paragraph = doc.add_paragraph()
                    doc.add_heading('SARA Recommendation', level=2)
                    doc.add_paragraph()       
                    paragraph = doc.add_paragraph(st.session_state["sara_recommendation_gpt"])

                    bio = io.BytesIO()
                    doc.save(bio)
                except NameError:
                    pass

            with col5_up:

                col_d1, col_d2 = st.tabs(["Download Report", "Download Case Package"])

                with col_d1:
                # Applying to download button -> download_button
                    st.markdown("""
                        <style>
                            .stButton download_button {
                                width: 100%;
                                height: 70%;
                            }
                        </style>
                    """, unsafe_allow_html=True)


                             
                    if doc:
                        st.download_button(
                            label="Download Report",
                            data=bio.getvalue(),
                            file_name="Report.docx",
                            mime="docx",
                            disabled=st.session_state.disabled
                        )
            with col_d2:
                
                # initiating a temp file
                tmp_dir = tempfile.mkdtemp()

                file_paths= []

                for uploaded_file in st.session_state.pdf_files:
                    file_pth = os.path.join(tmp_dir, uploaded_file.name)
                    with open(file_pth, "wb") as file_opn:
                        file_opn.write(uploaded_file.getbuffer())
                    file_paths.append(file_pth)

                for fetched_pdf in fetched_files:
                    # st.write(fetched_pdf)
                    file_pth = os.path.join('data/', fetched_pdf)
                    # st.write(file_pth)
                    file_paths.append(file_pth)

                
                combined_doc_path = os.path.join(tmp_dir, "report.docx")
                doc.save(combined_doc_path)



                # Create a zip file with the uploaded PDF files and the combined document
                zip_file_name = "package_files.zip"
                if file_paths:
                    files =  [combined_doc_path] + file_paths
                    create_zip_file(files, zip_file_name)
                    # create_zip_file(file_paths, zip_file_name)
                else:
                    pass

                
                # Download the package

                with open(zip_file_name, "rb") as file:
                    st.download_button(
                        label="Download Case Package", 
                        data=file, 
                        file_name=zip_file_name,
                        disabled=st.session_state.disabled)

                    # # Cleanup: Delete the temporary directory and its contents
                    # for file_path in file_paths + [combined_doc_path]:
                    #     os.remove(file_path)
                    # os.rmdir(temp_dir)

                with col6_up:   
                    # Adding Radio button
                    # st.markdown("""<span style="font-size: 24px; ">Make Decision</span>""", unsafe_allow_html=True)
                    if st.session_state.llm == "Closed-Source":
            
                        st.markdown("""<span style="font-size: 24px;color:#0000FF">Is SAR filing required?</span>""", unsafe_allow_html=True)

                        st.write("#### *SARA Recommendation*")
                        st.markdown("""<span style="font-size: 18px;">*Based on the following findings for the underlying case, under Bank Secrecy Act, it is recommended to file this case as a suspicious activity:*</span>""", unsafe_allow_html=True)
                        st.markdown("""<span style="font-size: 18px;">*1. Transaction amount is above the $5,000 value threshold*</span>""", unsafe_allow_html=True)
                        st.markdown("""<span style="font-size: 18px;">*2. There is an indication of suspicion with involvement of multiple individuals, mismatch of customer details on merchant invoice and identification of a potential suspect*.</span>""", unsafe_allow_html=True)  

                        st.warning('Please carefully review the recommendation and case details before the final submission',icon="⚠️")         
                            
                    if st.session_state.llm == "Open-Source":

                        st.write("#### *SARA Recommendation*")
                        st.markdown("""<span style="font-size: 18px;">*Based on the following findings,it is recommended to file this case as a suspicious activity to FinCEN under Bank Secrecy Act:*</span>""", unsafe_allow_html=True)
                        st.markdown("""<span style="font-size: 18px;">*1. Transaction amount is $5,600 indicating a need to file SAR with FinCEN.*</span>""", unsafe_allow_html=True)
                        st.markdown("""<span style="font-size: 18px;">*2. There is an indication of suspicion with identification of a suspect whose details mismatch with customer details on merchant invoice.*</span>""", unsafe_allow_html=True)  

                        st.warning('Please carefully review the recommendation and case details before the final submission',icon="⚠️")         
                    
                    
                    selected_rad = st.radio(":blue", ["opt1","Yes", "No", "Refer for review"], horizontal=True,disabled=st.session_state.disabled)
                    if selected_rad == "Refer for review":
                        email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
                        email_id = st.text_input("Enter email ID")
                        if email_id and not re.match(email_regex, email_id):
                            st.error("Please enter a valid email ID")


                    if st.button("Submit"):
                        if selected_rad in ("Yes"):
                            st.warning("Thanks for your review, your response has been submitted")
                        elif selected_rad in ("No"):
                            st.success("Thanks for your review, your response has been submitted")

                        else:
                            st.info("Thanks for your review, Case has been assigned to the next reviewer")






### AML code started

elif selected_option_case_type == "Money Laundering":
    st.markdown("### :blue[Anti-Money Laundering]")
    
# st.markdown('---')

    # Redirect to Merge PDFs page when "Merge PDFs" is selected
    if selected_option == "SAR-2023-24680":
        st.session_state.case_num = "SAR-2023-24680"
        
    
        col1,col2 = st.columns(2)
        # Row 1
        with col1:
            st.markdown("##### **Case number&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** SAR-2023-24680")
            st.markdown("##### **Customer name  :** Sarah Jones")
    
    
        with col2:
            st.markdown("##### **Case open date&nbsp;&nbsp;&nbsp;&nbsp;:** July 05, 2022")
            st.markdown("##### **Case type&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Money Laundering")
    
    
        # Row 2
        with col1:
            st.markdown("##### **Customer ID&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** 560062")
    
    
        with col2:
            st.markdown("##### **Case Status&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Open")

        st.markdown("---")
        
        ## Defining some global varibales for AML
        # directoty_path="ml_doc/"
        # fetched_files = read_pdf_files(directoty_path)
        directoty_path = "aml_docs/"
        fetched_files = read_pdf_files(directoty_path)
        
        


    
    


        if selected_option:
            col1_up, col2_up, col3_up, col4_up, col5_up, col6_up = st.tabs(["Data", "Generate Insights","Lineage","Summarization","Download Report", "Make a Decision"])

            with col1_up:
                bt1_up, bt2_up = st.tabs(["Fetch Evidence", "Upload Evidence"])

                with bt1_up:
                # Set the color
                # st.markdown(
                #     """
                #     <div style="display: flex; justify-content: center; align-items: center; height: 48px; border: 1px solid #ccc; border-radius: 5px; background-color: #f2f2f2;">
                #         <span style="font-size: 16px;  ">Fetch Evidence</span>
                #     </div>
                #     """,
                #     unsafe_allow_html=True
                # )
                    if 'clicked' not in st.session_state:
                        st.session_state.clicked = False
            
                    def set_clicked():
                        st.session_state.clicked = True
                        st.session_state.disabled = True
                    st.write("") #for the gap

                    st.button('Fetch Evidence', on_click=set_clicked)

                    if st.session_state.clicked:

                        # st.write("Evidence Files:") 
                        # st.markdown(html_str, unsafe_allow_html=True)
                
                        # Showing files
                        # show_files = fetched_files.copy()
                        # show_files = show_files + ['Other.pdf']
                        # files_frame = pd.DataFrame(show_files, columns=["File Name"])
                        # # files_frame["Select"] = [True for _ in range(len(files_frame))]
                        # files_frame = files_frame.reset_index(drop=True)

                        # # Add checkboxes to the DataFrame
                        # df_with_checkboxes = add_checkboxes_to_dataframe(files_frame)
                        
                        # # Iterate through each row and add checkboxes
                        # for index, row in df_with_checkboxes.iterrows():
                        #     if index < len(df_with_checkboxes) - 1:
                        #         checkbox_state = st.checkbox(f" {row['File Name']}", value=True)
                        #         df_with_checkboxes.loc[index, 'Select'] = checkbox_state
                        #     else:
                        #         st.checkbox(f"{row['File Name']}", value=False)



                        # st.dataframe(files_frame)
                        # st.write(df_reset.to_html(index=False), unsafe_allow_html=True)
                        # st.markdown(files_frame.style.hide(axis="index").to_html(), unsafe_allow_html=True)
                    
                        #select box to select file
                        fetch=["savings_account_statement.xlsx","credit_card_statement.xlsx","Transaction_Information.pdf","Customer_Relationship_Information.pdf","Blacklisted_Accounts.pdf"]
                        selected_file_name = st.selectbox(":blue[Select a file to View]",fetch)
                        st.write("Selected File: ", selected_file_name)
                        st.session_state.disabled = False
                        file_ext = tuple("pdf")  
                        file_ext1 = tuple("xlsx")
                    
                        if selected_file_name.endswith(file_ext):
                            selected_file_path = os.path.join(directoty_path, selected_file_name)
                            #converting pdf data to bytes so that render_pdf_as_images could read it
                            file = pdf_to_bytes(selected_file_path)
                            pdf_images = render_pdf_as_images(file)
                            #showing content of the pdf
                            st.subheader(f"Contents of {selected_file_name}")
                            for img_bytes in pdf_images:
                                st.image(img_bytes, use_column_width=True)
                        
                        elif selected_file_name.endswith(file_ext1):
                            selected_file_path = os.path.join(directoty_path, selected_file_name)
                            # from openpyxl import reader,load_workbook,Workbook
                            # wb=Workbook()
                            # wb = load_workbook(selected_file_path, read_only=False)
                            # st.write(wb.sheetnames)
                            # st.title(wb)
                            # st.write(wb.active)
                            # Read the Excel file into a DataFrame
                            df = pd.read_excel(selected_file_path, engine='openpyxl')
                            # Find the row index where the table data starts
                            data_start_row = 0  # Initialize to 0
                            for i, row in df.iterrows():
                                if row.notna().all():
                                    data_start_row = i
                                    break       
                            if data_start_row>0:  
                                df.columns = df.iloc[data_start_row]
                            # Extract the text content above the data
                            text_content = "\n".join(df.iloc[:data_start_row].apply(lambda x: "\t".join(map(str, x)), axis=1)).replace('nan','')
                            
                            
                            st.text(text_content)
                            
                            st.write(df.iloc[data_start_row+1:], index=False)
                    
                    
                        else:
                            selected_file_path = os.path.join(directoty_path, selected_file_name)
                            # This is showing png,jpeg files
                            st.image(selected_file_path, use_column_width=True)
        
    
    
                with bt2_up:
                    pdf_files = st.file_uploader("", type=["pdf","png","jpeg","docx","xlsx"], accept_multiple_files=True)
                    st.session_state.pdf_files = pdf_files
                    # showing files
                    for uploaded_file in pdf_files:
                        #This code is to show pdf files
                        file_ext = tuple("pdf")
                        if uploaded_file.name.endswith(file_ext):
                            # Show uploaded files in a dropdown
                            # if pdf_files:
                            st.subheader("Uploaded Files")
                            file_names = [file.name for file in pdf_files]
                            selected_file = st.selectbox(":blue[Select a file]", file_names)
                            # Enabling the button
                            st.session_state.disabled = False
                            # Display selected PDF contents
                            if selected_file:
                                selected_pdf = [pdf for pdf in pdf_files if pdf.name == selected_file][0]
                                pdf_images = render_pdf_as_images(selected_pdf)
                                st.subheader(f"Contents of {selected_file}")
                                for img_bytes in pdf_images:
                                    st.image(img_bytes, use_column_width=True)
        
                        else:
                            # This is showing png,jpeg files
                            st.image(uploaded_file, use_column_width=True)
        
                #creating temp directory to have all the files at one place for accessing
                tmp_dir_ = tempfile.mkdtemp()
                temp_file_path1= []
                temp_file_path2= []
    
                for uploaded_file in pdf_files:
                    file_ext = tuple("pdf")
                    if uploaded_file.name.endswith(file_ext):
                        file_pth = os.path.join(tmp_dir_, uploaded_file.name)
                        with open(file_pth, "wb") as file_opn:
                            file_opn.write(uploaded_file.getbuffer())
                            temp_file_path2.append(file_pth)
                    else:
                        pass
    
    
                for fetched_pdf in fetched_files:
                    directory_path="aml_docs/"
        
                    
                    file_ext1 = tuple("pdf")
                    file_ext2 = tuple(["xlsx","csv"])
                    file = fetched_pdf.split('.',1)[0]
                    if fetched_pdf.endswith(file_ext1):
                        selected_file_path = os.path.join(directory_path, fetched_pdf)
                        if is_searchable_pdf(selected_file_path)==False:
                            text = convert_scanned_pdf_to_searchable_pdf(selected_file_path)
                            texts =  text_to_docs(text,file)
                            for i in texts:
                                temp_file_path2.append(i)
                        else:
                            file_pth = os.path.join(directory_path, fetched_pdf)
                            text = extract_text_from_pdf(file_pth)
                            # st.write(text)
                            texts =  text_to_docs(text,file)
                            for i in texts:
                                temp_file_path2.append(i)
                    elif fetched_pdf.endswith(file_ext2):
                        selected_file_path = os.path.join(directory_path, fetched_pdf)
                        
                        if selected_file_path.startswith("aml_docs/credit_card_statement"):
                            
                            json1=process_data_credit_card(selected_file_path)
                            #st.write("creditcard")
                            #st.write(json1)
                            #text = convert_image_to_searchable_pdf(selected_file_path)
                            texts = text_to_docs(json1,file)
                            for i in texts:
                                temp_file_path2.append(i)
                        elif selected_file_path.startswith("aml_docs/savings_account_statement"):
                            #selected_file_path = os.path.join(directory_path, fetched_pdf)
                            json2=process_data_saving(selected_file_path)
                            #st.write("savings")
                            #st.write(json2)
                            #text = convert_image_to_searchable_pdf(selected_file_path)
                            texts = text_to_docs(json2,file)
                            for i in texts:
                                temp_file_path2.append(i)
                #st.write(temp_file_path2)            

    
                #combining files in fetch evidence and upload evidence
                


        

            with col2_up:
                #This is the embedding model
                model_name = "thenlper/gte-small"
                # model_name = "hkunlp/instructor-large"
                
                # Memory setup for gpt-3.5
                llm = ChatOpenAI(temperature=0.1)
                memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=200)
                conversation = ConversationChain(llm=llm, memory =memory,verbose=False)
                
        
                # Adding condition on embedding
                try:
                    if temp_file_path2:
                        hf_embeddings = embed(model_name) 
                    else:
                        pass
                except NameError:
                    pass
                
                # Chunking with overlap
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size = 700,
                    chunk_overlap  = 0,
                    length_function = len,
                    separators=["\n\n", "\n", " ", ""]
                )
                # try:
                #     if temp_file_path2:
                #         docs, docsearch = embedding_store(temp_file_path2)
                #     else:
                #         pass
                # except Exception:
                #     pass            

                # Creating header
                col1,col2 = st.columns(2)
                with col1:
                    st.markdown("""<span style="font-size: 24px; ">Key Questions</span>""", unsafe_allow_html=True)
        
                    # Create a Pandas DataFrame with your data

                    data = {'Questions': ["Why was the transaction triggered?",
                          "What are the products that are associated with this customer?",
                         "What are the associated suspicious transactions for Credit Card?",
                         # "What is the total amount associated with the money laundering activity for Credit card?",
                          "What are the associated suspicious transactions for Savings account?",
                         # "What is the total amount associated with the money laundering activity for Savings Account?",
                          #"What type of Money laundering activity is taking place?",
                          "What is the total amount associated with the Money Laundering ?"]}
            
                    df_fixed = pd.DataFrame(data)
                    df_fixed.index = df_fixed.index +1
                with col2:

                    # Create a checkbox to show/hide the table
                    cols1, cols2, cols3, cols4 = st.columns([1,1,1,1])
                    with cols1:
                        show_table2 = tog.st_toggle_switch(label="", 
                                            key="Key1", 
                                            default_value=False, 
                                            label_after = False, 
                                            inactive_color = '#D3D3D3', 
                                            active_color="#11567f", 
                                            track_color="#29B5E8"
                                            )
                    # Show the table if the checkbox is ticked
                    if show_table2:
                        # st.write(df_fixed)
                        # st.dataframe(df_fixed, width=1000)
                        df_fixed["S.No."] = df_fixed.index
                        df_fixed = df_fixed.loc[:,['S.No.','Questions']]
                        st.markdown(df_fixed.style.hide(axis="index").to_html(), unsafe_allow_html=True)
                
                    
                with st.spinner('Wait for it...'):
                    
                    if 'clicked1' not in st.session_state:
                        st.session_state.clicked1 = False
                   
                    def set_clicked1():
                        st.session_state.clicked1 = True
                        st.session_state.disabled = True
 
                   
                    generate_button = st.button("Generate Insights",on_click=set_clicked1,disabled=st.session_state.disabled)
                   
 
                    if st.session_state.clicked1:
                        if temp_file_path2 is not None:
                            
                            doc_1, docsearch2 = embedding_store_aml_2(temp_file_path2,hf_embeddings)
                            # File handling logic
                            
                            
                           

                            
                            if st.session_state.llm == "Closed-Source":
                                chat_history_1 = {}

                                ## Question-1

                                
                                # st.write(doc_1)
                                # st.write(temp_file_path2)

    
                                query = "Why was the transaction triggered?"
                                
                                context_1 = docsearch2.similarity_search(query, k=5)
                                #st.write(context_1)
                                prompt_1 = f'''You should closely look into the transactions information data for the reason why was the transaction flagged as suspicious. \n\n
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: Give a concise response as reason in one sentence. '''
                                #response = usellm(prompt_1)
                                system_prompt = wrap_prompt("You are a Money Laundering Analyst.", "system")
                                user_prompt = wrap_prompt(prompt_1, "user")
                                res = get_response([system_prompt, user_prompt])
                                response = res['choices'][0]['message']['content']
                                ques1 = response
                                
                                
                                
                                
                                chat_history_1[query] = response
                                st.session_state["lineage_aml"][query] = context_1
                               
    
                                ## Question-2

                                
                                query = "What are the products that are associated with this customer?"
                                context_1 = docsearch2.similarity_search(query, k=5)
                                prompt_1 = f'''Your goal is identify all the products that are associated with the customer. \n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Output the identified Products, Do not give/add any Explanation, Note, etc. in the answer.)'''
                                #response = usellm(prompt_1)
                                system_prompt = wrap_prompt("You are a Money Laundering Analyst.", "system")
                                user_prompt = wrap_prompt(prompt_1, "user")
                                res = get_response([system_prompt, user_prompt])
                                response = res['choices'][0]['message']['content']
                                
                                
                                chat_history_1[query] = response
                                st.session_state["lineage_aml"][query] = context_1

                                

                                ## Question-3

                                query = "What are the associated suspicious transactions for Credit Card?"
                                context_1 = docsearch2.similarity_search(query, k=5)
                                prompt_1=f''' Your goal is to identify the suspicious transactions from Credit_Card_statement. Suspicious transactions can be:\n\n
                                Transactions made to a suspicious entity. Output "Description", "Date" and "Debited ($)" of those identified transactions. # Strictly do not repeat any transaction.\n\
                                Context: {context_1}\n\
                                Response: (Do not give/add any extra Note, Explanation in answer.) '''
                                
                                #st.write(context_1)
                                system_prompt = wrap_prompt("You are a Money Laundering Analyst.", "system")
                                user_prompt = wrap_prompt(prompt_1, "user")
                                res = get_response([system_prompt, user_prompt])
                                response = res['choices'][0]['message']['content']
                                transactions_cc = response
                                chat_history_1[query] = response
                                st.session_state["lineage_aml"][query] = context_1

                                ## Question-3.1

                                query = "What is the total amount associated with the money laundering activity for Credit card?"
                                #st.session_state["lineage_aml"][query] = context_1
                
                                context_1 = transactions_cc
                                prompt_1 = f'''Act as a calculator and add up all the transactions amount in the context.\n\
                                Output the total calculated amount as answer to the question.
                                Context: {context_1}\n\
                                Question: {query}\n\
                                Response: (Add this before the total amount : "Total Money Laundering amount that can be associated with credit card is : ")'''


                                system_prompt = wrap_prompt("You are a Money Laundering Analyst.", "system")
                                user_prompt = wrap_prompt(prompt_1, "user")
                                res = get_response([system_prompt, user_prompt])
                                response = res['choices'][0]['message']['content']
                                #response = response.replace("33000", "USD 33000")
                                response = response.replace("$", "USD ")
                                total_cc = response

                                ## Question-4

             

                                query = "What are the associated suspicious transactions for Savings account?"
                                context_1 = docsearch2.similarity_search(query, k=5)
                                  

                                prompt_1=f''' Your goal is to identify the suspicious transactions from savings_account_statement. Suspicious transactions can be:\n\n
                                High Value Cash Deposits in a short span of time. Strictly do not include any Paycheck transactions and Opening balance transaction as they may not be considered as suspicious transactions. Output the "Description", "Date" and "Credited ($)" of those identified transactions.Also, do not repeat the same transaction.\n\
                                Context: {context_1}\n\
                                Response: (Strictly do not give/add any Note, Explanation in answer.) '''
                                #st.write(context_1)
                                system_prompt = wrap_prompt("You are a Money Laundering Analyst.", "system")
                                user_prompt = wrap_prompt(prompt_1, "user")
                                res = get_response([system_prompt, user_prompt])
                                response = res['choices'][0]['message']['content']
                                #response = usellm(prompt_1)

                               
                                transactions_sa = response
                                
                                chat_history_1[query] = response
                                st.session_state["lineage_aml"][query] = context_1
                                context_q5 = context_1

                                ## Question-4.1

                                query = "What is the total amount associated with the money laundering activity for Savings Account ?"
                                #st.session_state["lineage_aml"][query] = context_1
                                context_1 = transactions_sa
                                prompt_1 = f'''Act as a calculator and add up all the transactions amount in the context.\n\
                                Output the total calculated amount as answer to the question.
                                Context: {context_1}\n\
                                Question: {query}\n\
                                Response: (Add this before the toal amount : "Total Money Laundering amount that can be associated with savings account is : ")'''
                                

                                response = usellm(prompt_1)
                                
                                #response = response.replace("33000", "USD 33000")
                                response = response.replace("$", "USD ")
                                total_sav =  response
                                
                                # ## Question-5.1

                                # query = "What type of Money laundering activity is taking place?"
                                # context_1 = docsearch.similarity_search(query, k=5)
                                  

                                # prompt_1=f'''You Are an Anti-Money Laundering Specialist, carefully observe the transaction statements pattern from both the transactions data of credit card and saving accounts statements combined. \
                                # The type of money laundering activities which can take place includes: Structuring or smurfing, layering, round tripping, etc.\ 
                                # Act as and Anti-Money Laundering analyst, observe the transactions statements data and give a concise answer with explanation of what type of money laundering activity could be taking place and on what pattern this activity is observed.\n\n
                                # Question: {query}\n\
                                # Context: {context_1}\n\
                                # Response: (Give me a concise response in one sentence stating the type of money laundering activity the can be taking place and on what patterns it is observed . Do not give me any Note etc)'''

                                # response = usellm(prompt_1)
                                
                                ## Question-5

                                query = "What is the total amount associated with the Money Laundering ?"
                                st.session_state["lineage_aml"][query] = context_q5
                                context_1 = transactions_cc + transactions_sa
                                  

                                prompt_1 = f'''Based on the Context, what is the relationship between the suspicious transactions of savings accounts and credit card transactions.\n\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence stating what TYPE of money laundering activity is taking place and WHY, along with the relationship found? Do not give me any Note etc).'''

                                #response = usellm(prompt_1)
                                system_prompt = wrap_prompt("You are a Money Laundering Analyst.", "system")
                                user_prompt = wrap_prompt(prompt_1, "user")
                                res = get_response([system_prompt, user_prompt])
                                response = res['choices'][0]['message']['content']
                                response1 = total_sav + " and "+ total_cc + "  ."+ response
                                ques5 = response1
                                chat_history_1[query] = response1
                                


                        
    
                                try:
                                    res_df_gpt = pd.DataFrame(list(chat_history_1.items()), columns=['Question','Answer'])
                                    res_df_gpt.reset_index(drop=True, inplace=True)
                                    index_ = pd.Series([1,2,3,4,5])
                                    res_df_gpt = res_df_gpt.set_index([index_])
                                    # st.write(res_df_gpt)                             
                                except: 
                                    e = Exception("")
                                    st.exception(e)

                            
                                
                                #Display table
                                st.table(res_df_gpt)

                                #copy in session state
                                st.session_state["tmp_table_gpt_aml"] = pd.concat([st.session_state.tmp_table_gpt_aml, res_df_gpt], ignore_index=True)
                                
                                ################## SARA Recommendation ######################

                                query  = "Give your recommendation if this is a Suspicious activity or not?"
                                contexts = ', '.join(res_df_gpt['Answer'])
                                prompt_2 = f"""Is this a case of Suspicious activity? If yes, then Find answer to the questions as truthfully as possible as per the available information only,\n\n\
                                1.) why was the transaction triggered?\n\
                                2.) what are the total amounts related to money laundering for savings account and credit cards?\n\
                                3.) what type of money laundering activity is taking place and why ?\n\n\                     
                                Context: {contexts}\n\
                                Also, add your concise recommendation whether SAR filling is required or not ?
                                Response: start the output answering if it can be considered as a suspicious activity or not based on the avaliable information in a sentence, then answer all the questions as individual points."""
                                system_prompt = wrap_prompt("You are a Money Laundering Analyst.", "system")
                                user_prompt = wrap_prompt(prompt_2, "user")
                                res = get_response([system_prompt, user_prompt])
                                response = res['choices'][0]['message']['content']
                                #response1 = usellm(prompt)
                                response1 = response.replace("$", "USD ")
                                sara_close_source=response1
              


                                st.session_state["sara_recommendation_gpt_aml"] = response1                
                                
                                st.markdown("### SARA Recommendation")
                                
                                st.markdown(response1)
                                

                                st.markdown("#### Recommendation Feedback:")
                                col_1, col_2, col_3, col_4, col_5, col_6 = st.columns(6)
                    
                                with col_1:
                                    if st.button("👍🏻",key=2):
                                        st.write("*Feedback is recorded*")
                    
                    
                                with col_2:
                                    if st.button("👎🏻",key=3):
                                        st.write("*Feedback is recorded*")


                            elif st.session_state.llm == "Open-Source":
    
                                chat_history = {}

                                ## question-1 
    
                                query = "Why was the transaction triggered?"
                                context_1 = docsearch2.similarity_search(query, k=5)
                                prompt_1 = f'''You should closely look into the transactions information data for the reason why was the transaction flagged as suspicious. \n\n
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: Give a concise response as reason in one sentence. '''
                                response = llama_llm(llama_13b,prompt_1)
                                question_1 = response
                                chat_history[query] = response
                                st.session_state["lineage_aml_llama"][query] = context_1

                                ##question-2
                  
    
                                query = "What are the products that are associsted with this customer?"
                                context_1 = docsearch2.similarity_search(query, k=5)
                                prompt_1 = f'''Your goal is identify all the products that are associated with the customer. \n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Output the identified Products, Do not give/add any Explanation, Note, etc. in the answer.)'''

                                response = llama_llm(llama_13b,prompt_1)
                                chat_history[query] = response
                                st.session_state["lineage_aml_llama"][query] = context_1

                            

                                #question-3
                            
                                query = "What are the associated suspicious transactions for Credit Card?"
                                context_1 = docsearch2.similarity_search(query, k=5)
                                prompt_1 = f''' Your goal is to identify the suspicious transactions from Credit_Card_statement. Suspicious transactions can be:\n\n
                                Transactions made to a suspicious entity. Output "Description", "Date" and "Debited ($)" of those identified transactions. # Strictly do not repeat any transaction.\n\
                                Context: {context_1}\n\
                                Response: (Do not give/add any extra Note, Explanation in answer.) '''

                                response = llama_llm(llama_13b,prompt_1)
                                transactions_cc_llama = response
                                chat_history[query] = response
                                st.session_state["lineage_aml_llama"][query] = context_1

                                ## question-3.1

                                query = "What is the total amount associated with the money laundering activity for Credit card?"
                                #st.session_state["lineage_aml_llama"][query] = context_1
                                context_1 = transactions_cc_llama
                                prompt_1 = f'''Act as a calculator and add up all the transactions amount in the context.\n\
                                Output the total calculated amount as answer to the question.
                                Context: {context_1}\n\
                                Question: {query}\n\
                                Response: (Add this before the total amount : "Total Money Laundering amount that can be associated with credit card is : ")'''
          
                                response = llama_llm(llama_13b,prompt_1)
                                response = response.replace("$", "USD ")
                                total_cc_llama = response
                                #chat_history[query] = response

                                ## question-4

                                query = "What are the associated suspicious transactions for Savings account?"
                                context_1 = docsearch2.similarity_search(query, k=5)
                                prompt_1 = f''' Your goal is to identify the suspicious transactions from savings_account_statement. Suspicious transactions can be:\n\n
                                High Value Cash Deposits in a short span of time. Strictly do not include any Paycheck transactions and Opening balance transaction as they may not be considered as suspicious transactions. Output the "Description", "Date" and "Credited ($)" of those identified transactions.Also, do not repeat the same transaction.\n\
                                Context: {context_1}\n\
                                Response: (Strictly do not give/add any Note, Explanation in answer.) '''

                                response = llama_llm(llama_13b,prompt_1)
                                transactions_sa_llama = response
                                chat_history[query] = response
                                st.session_state["lineage_aml_llama"][query] = context_1

                                ## question-4.1

                                query = "What is the total amount associated with the money laundering activity for Savings Account ?"
                                #st.session_state["lineage_aml_llama"][query] = context_1
                                context_1 = transactions_sa_llama
                                prompt_1 = f'''Act as a calculator and add up all the transactions amount in the context.\n\
                                Output the total calculated amount as answer to the question.
                                Context: {context_1}\n\
                                Question: {query}\n\
                                Response: (Add this before the toal amount : "Total Money Laundering amount that can be associated with savings account is : ")'''
          
                                response = llama_llm(llama_13b,prompt_1)
                                response = response.replace("$", "USD ")
                                total_sav_llama = response
                                #chat_history[query] = response

                                ## question-5.1

                                query = "What type of Money laundering activity is taking place?"
                                context_1 = docsearch2.similarity_search(query, k=5)
                                prompt_1 = f'''You Are an Anti-Money Laundering Specialist, carefully observe the transaction statements pattern from both the transactions data of credit card and saving accounts statements combined. \
                                The type of money laundering activities which can take place includes: Structuring or smurfing, layering, round tripping, etc.\ 
                                Act as and Anti-Money Laundering analyst, observe the transactions statements data and give a concise answer with explanation of what type of money laundering activity could be taking place and on what pattern this activity is observed.\n\n
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence stating the type of money laundering activity the can be taking place and on what patterns it is observed . Do not give me any Note etc)'''

                                response = llama_llm(llama_13b,prompt_1)
                                
                                # chat_history[query] = response
                                # st.session_state["lineage_aml_llama"][query] = context_1

                                ## question-5

                                query = "What is the total amount associated with the Money Laundering ?"
                                st.session_state["lineage_aml_llama"][query] = context_1
                                context_1 = transactions_cc_llama + transactions_sa_llama
                                prompt_1 = f'''Based on the Context, what is the relationship between the suspicious transactions of savings accounts and credit card transactions.\n\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence stating the type of money laundering activity the can be taking place and on what patterns it is observed along with the relationship found. Do not give me any Note etc)'''

                                response = llama_llm(llama_13b,prompt_1)
                                response = total_sav_llama + " and "+ total_cc_llama + "  ."+ response
                                question_8 = response
                                chat_history[query] = response



                            
                   
                                try:
                                    res_df_llama = pd.DataFrame(list(chat_history.items()), columns=['Question','Answer'])
                                    res_df_llama.reset_index(drop=True, inplace=True)
                                    index_ = pd.Series([1,2,3,4])
                                    res_df_llama = res_df_llama.set_index([index_])
                                    # st.write(res_df_llama)
                            
                                except: 
                                    e = Exception("")
                                    st.exception(e)

                                st.table(res_df_llama)
                                #copy in session state
                                st.session_state["tmp_table_llama_aml"] = pd.concat([st.session_state.tmp_table_llama_aml, res_df_llama], ignore_index=True)

                                
                                ################## SARA Recommendation ######################


                    
                    
                                queries ="Give your recommendation if this is a Suspicious activity or not?"
                    
                                contexts = question_1 + question_8
                                prompt = f"""Give concise response to the each questions below within the given Context. \n\
                                1.) transaction triggered\n\
                                2.) amounts related to money laundering for savings account and credit cards\n\
                                3.) Type of money laundering activity taking place and why ?\n\                          
                                Context: {contexts}\n\
                                Response: (Give a neatly formatted response for each question individually. Also, give your recommendation for the below Question.) 
                                Question: {queries} """
                                                    
                                response1 = llama_llm(llama_13b,prompt)    
                                
                                response1 = response1.replace("$", "USD ") 
                                sara_open_source = response1      
                                
                                
                                st.session_state["sara_recommendation_llama_aml"] = response1                    

                                st.markdown("### SARA Recommendation")
                                st.markdown(response1)

                                st.markdown("#### Recommendation Feedback:")
                                col_1, col_2, col_3, col_4, col_5, col_6 = st.columns(6)
                    
                                with col_1:
                                    if st.button("👍🏻",key=2):
                                        st.write("*Feedback is recorded*")
                    
                    
                                with col_2:
                                    if st.button("👎🏻",key=3):
                                        st.write("*Feedback is recorded*")
                                
                              
                
    
    
                st.markdown("---")
    
                # For input box outside of template4
                try:
                    if temp_file_path2:
                        docs, docsearch = embedding_store_aml(temp_file_path2)
                    else:
                        pass
                except Exception:
                    pass
    
    
                # Text Input
                # st.markdown("""<span style="font-size: 24px; ">Ask Additional Questions</span>""", unsafe_allow_html=True)
                query = st.text_input(':black[Ask Additional Questions]',disabled=st.session_state.disabled)
                text_dict = {}
                @st.cache_data
                def LLM_Response():
                    llm_chain = LLMChain(prompt=prompt, llm=llm)
                    response = llm_chain.run({"query":query, "context":context})
                    return response
                if st.session_state.llm == "Closed-Source":
                    with st.spinner('Getting you information...'):      
                        if query:
                            # docs = chunk_extract(temp_file_path2)
                            # text_data_doc = context_data(docs)
                            
                            # Text input handling logic
                            #st.write("Text Input:")
                            #st.write(text_input)
                
                            context_1 = docsearch.similarity_search(query, k=5)
                            st.session_state.context_1 = context_1
                            
                        
                            prompt_1 = f''' You Are an Anti-Money Laundering Specialist, provide the answer to the below question in a concise manner.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\                      
                                            Response: '''
                                                
                            
    
                            #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                            response = usellm(prompt_1) #LLM_Response()
                            text_dict[query] = response
                            #Display response
                            st.write(response)
                            if response:
                                df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                            else:
                                df = pd.DataFrame()
                
                            st.session_state["tmp_table_gpt_aml"] = pd.concat([st.session_state.tmp_table_gpt_aml, df], ignore_index=True)
                            st.session_state.tmp_table_gpt_aml.drop_duplicates(subset=['Question'])
                
                
                elif st.session_state.llm == "Open-Source":
                    with st.spinner('Getting you information...'):      
                        if query:
                            # docs = chunk_extract(temp_file_path2)
                            # text_data_doc = docs
                            #text_data_doc = process_files_and_generate_responses(fetched_files)
                            # Text input handling logic
                            #st.write("Text Input:")
                            #st.write(text_input)
                
                            context_1 = docsearch.similarity_search(query, k=5)
                            st.session_state.context_1 = context_1
                            prompt_1 = f''' You Are an Anti-Money Laundering Specialist, provide the answer to the below question in a concise manner.\n\n\
                                            Question: {query}\n\
                                            Context: {context_1}\n\                      
                                            Response: '''
                    
                            #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                            # response = usellm(prompt_1) #LLM_Response()
                            response = llama_llm(llama_13b,prompt_1)
                            text_dict[query] = response
                
                            st.write(response)
                
                            if response:
                                df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                            else:
                                df = pd.DataFrame()
                
                            st.session_state["tmp_table_llama_aml"] = pd.concat([st.session_state.tmp_table_llama_aml, df], ignore_index=True)
                            
                            st.session_state.tmp_table_llama_aml.drop_duplicates(subset=['Question'])
                
            with col3_up:
                if st.session_state["lineage_aml"] is not None:
                    
                    
 
                    li = ["Select question to get the lineage",
                        "Why was the transaction triggered?",
                        "What are the products that are associsted with this customer?",
                        "What are the associated suspicious transactions for Credit Card?",
                        #"What is the total amount associated with the money laundering activity for Credit card?",
                        "What are the associated suspicious transactions for Savings account?",
                        #"What is the total amount associated with the money laundering activity for Savings Account ?",
                        #"What type of Money laundering activity is taking place?",
                        "What is the total amount associated with the Money Laundering ?",
                        ]
                    
                   
                    selected_option = st.selectbox("", li)
                    if selected_option in li[1:]:

        
                        doc = st.session_state["lineage_aml"][selected_option]
                        
                        
                        for i in range(len(doc)):
                            #st.write(doc[i])
                            y=i+1
                            st.write(f":blue[Chunk-{y}:]")
                            st_ = doc[i].page_content.replace("($)"," ")
                            st.write(":blue[Page Content:]",st_) 
                            st.write(":blue[Source:]",doc[i].metadata['source'])
                              

  
            
                    

            with col4_up:
                with st.spinner('Summarization ...'):
                    st.markdown("""<span style="font-size: 24px; ">Summarize key findings of the case.</span>""", unsafe_allow_html=True)
                    st.write()
                    if st.button("Summarize",disabled=st.session_state.disabled):
                        if st.session_state.llm == "Closed-Source":
                            st.session_state.disabled=False
            
                            # summ_dict_gpt = st.session_state.tmp_table_gpt_aml.set_index('Question')['Answer'].to_dict()
                            summary1= ', '.join(res_df_gpt['Answer']) + sara_close_source
                            # chat_history = resp_dict_obj['Summary']
                            # memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=400)
                            # memory.save_context({"input": "This is the entire summary"}, 
                            #                 {"output": f"{summary1}"})
                            # conversation = ConversationChain(
                            # llm=llm, 
                            # memory = memory,
                            # verbose=True)
                            # st.write(summ_dict_gpt)
                            # st.write(summary1)
                            # st.session_state["tmp_summary_gpt_aml"] = conversation.predict(
                            #     input="Act as a summarization tool and Provide a detailed summary of the provided information include all the relevant information and numbers. Provide the summary in a single paragraph and don't include words like these: 'chat summary', 'includes information' or 'AI' in my final summary.")
                            # st.session_state["tmp_summary_gpt_aml"]=st.session_state["tmp_summary_gpt_aml"].replace("$", "USD ")

                            ## using open ai:

                            prompt_summ=f'''Provide a detailed summary of the below Context, and make sure to include all the relevant information (this includes names, transactions, involved parties, amounts involved, etc.). Provide the summary in a single paragraph and don't include words like these: 'chat summary', 'includes information' or 'AI' in my final summary.\n\n\
                            Context: {summary1}  '''
                            system_prompt = wrap_prompt("You are a summarization tool", "system")
                            user_prompt = wrap_prompt(prompt_summ, "user")
                            res = get_response([system_prompt, user_prompt])
                            response = res['choices'][0]['message']['content']
                            response_summary = response.replace("$", "USD ")

                            
                            st.session_state["tmp_summary_gpt_aml"]=response_summary
                            
                            #Display summary
                            st.write(st.session_state["tmp_summary_gpt_aml"])

                            st.markdown("#### Summarization Feedback:")
                            col_1, col_2, col_3, col_4, col_5, col_6 = st.columns(6)
                    
                            with col_1:
                                if st.button("👍🏻",key=4):
                                    st.write("*Feedback is recorded*")
                
                
                            with col_2:
                                if st.button("👎🏻",key=5):
                                    st.write("*Feedback is recorded*")


                        elif st.session_state.llm == "Open-Source":
                            st.session_state.disabled=False
                            template = """ You are a summarization tool and your goal is to summarize the data provided in such a way that it includes all the relevant information. # DO not use words such as AI or tool
                            ```{text}```
                            Response: """
                            prompt = PromptTemplate(template=template,input_variables=["text"])
                            llm_chain_llama = LLMChain(prompt=prompt,llm=llama_13b)

                            summ_dict_llama = st.session_state.tmp_table_llama_aml.set_index('Question')['Answer']
                            text = []
                            for key,value in summ_dict_llama.items():
                                text.append(value)
                            st.session_state["tmp_summary_llama_aml"] = llm_chain_llama.run(text)
                            st.session_state["tmp_summary_llama_aml"]=st.session_state["tmp_summary_llama_aml"].replace("$", "USD ")
                            #Display summary
                            st.write(st.session_state["tmp_summary_llama_aml"])

                            st.markdown("#### Summarization Feedback:")
                            col_1, col_2, col_3, col_4, col_5, col_6 = st.columns(6)
                    
                            with col_1:
                                if st.button("👍🏻",key=4):
                                    st.write("*Feedback is recorded*")
                
                
                            with col_2:
                                if st.button("👎🏻",key=5):
                                    st.write("*Feedback is recorded*")

        
            tmp_summary = []
            tmp_table = pd.DataFrame()

            try: 
                if st.session_state.llm == "Closed-Source":
                    st.session_state.disabled=False
                    tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_gpt_aml"]], ignore_index=True)
                    tmp_summary.append(st.session_state["tmp_summary_gpt_aml"])
                    tmp_table.drop_duplicates(inplace=True)
                    #st.write(tmp_table)


                elif st.session_state.llm == "Open-Source":
                    st.session_state.disabled=False
                    tmp_summary.append(st.session_state["tmp_summary_llama_aml"])
                    tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_llama_aml"]], ignore_index=True)
                    tmp_table.drop_duplicates(inplace=True)

            except:
                e = Exception("")
                st.exception(e)


            try:
                # initiate the doc file
                doc = docx.Document()
                # doc.add_section(WD_SECTION.NEW_PAGE)
                doc.add_heading(f"Case No.: {st.session_state.case_num}",0)

                # Add a subheader for case details
                subheader_case = doc.add_paragraph("Case Details")
                subheader_case.style = "Heading 2"
                # Addition of case details
                paragraph = doc.add_paragraph(" ")
                case_info = {
                    "Case Number                            ": " SAR-2023-24680",
                    "Customer Name                       ": " Sarah Jones",
                    "Customer ID                              ": " 560062",
                    "Case open date                         ": " July 05, 2022",
                    "Case Type                                  ": " Money Laundering",
                    "Case Status                                ": " Open"
                }
                for key_c, value_c in case_info.items():
                    doc.add_paragraph(f"{key_c}: {value_c}")
                paragraph = doc.add_paragraph(" ")

                # Add a subheader for customer info to the document ->>
                subheader_paragraph = doc.add_paragraph("Customer Information")
                subheader_paragraph.style = "Heading 2"
                paragraph = doc.add_paragraph(" ")

                # Add the customer information
                customer_info = {
                    "Name                                           ": " Sarah Jones",
                    "Address                                      ": " 858 3rd Ave, Chula Vista, California, 91911 US",
                    "Phone                                          ": " (619) 425-2972",
                    "A/C No.                                        ": " 4587236908230087",
                    "SSN                                               ": " 653-30-9562"
                }

                for key, value in customer_info.items():
                    doc.add_paragraph(f"{key}: {value}")
                paragraph = doc.add_paragraph()
                # Add a subheader for Suspect infor to the document ->>
                subheader_paragraph = doc.add_paragraph("Suspect's Info")
                subheader_paragraph.style = "Heading 2"
                paragraph = doc.add_paragraph()
                #""" Addition of a checkbox where unticked box imply unavailability of suspect info"""
                
                doc.add_heading('Summary', level=2)
                paragraph = doc.add_paragraph()
                doc.add_paragraph(tmp_summary)
                paragraph = doc.add_paragraph()
                doc.add_heading('Key Insights', level=2)
                paragraph = doc.add_paragraph()
                columns = list(tmp_table.columns)
                table = doc.add_table(rows=1, cols=len(columns), style="Table Grid")
                table.autofit = True
                
                for col in range(len(columns)):
                    # set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50) # set cell margin
                    table.cell(0, col).text = columns[col]
                # doc.add_table(st.session_state.tmp_table.shape[0]+1, st.session_state.tmp_table.shape[1], style='Table Grid')

                for i, row in enumerate(tmp_table.itertuples()):
                    table_row = table.add_row().cells # add new row to table
                    for col in range(len(columns)): # iterate over each column in row and add text
                        table_row[col].text = str(row[col+1]) # avoid index by adding col+1
                # save document
                # output_bytes = docx.Document.save(doc, 'output.docx')
                # st.download_button(label='Download Report', data=output_bytes, file_name='evidence.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                paragraph = doc.add_paragraph()
                paragraph = doc.add_paragraph()
                doc.add_heading('SARA Recommendation', level=2)
                doc.add_paragraph()       
                paragraph = doc.add_paragraph(st.session_state["sara_recommendation_gpt"])

                bio = io.BytesIO()
                doc.save(bio)
            except NameError:
                pass
            
            
            with col5_up:

                col_d1, col_d2 = st.tabs(["Download Report", "Download Case Package"])

                with col_d1:
                    # Applying to download button -> download_button
                    st.markdown("""
                        <style>
                            .stButton download_button {
                                width: 100%;
                                height: 70%;
                            }
                        </style>
                    """, unsafe_allow_html=True)


                    # combined_doc_path = os.path.join(tmp_dir, "resulting_document.docx")
                    # doc.save(combined_doc_path)

                    # # Create a zip file with the uploaded PDF files and the combined document
                    # zip_file_name = "package_files.zip"
                    # if pdf_files:
                    #     st.write(file_paths)
                    #     files =  [combined_doc_path]
                    #     st.write(files)
                        
                    #     create_zip_file(files, zip_file_name)
                    #     # create_zip_file(file_paths, zip_file_name)
                    # else:
                    #     pass
                    # # Download the package
                    # with open(zip_file_name, "rb") as file:
                    #     st.download_button(
                    #         label="Download Case Package", 
                    #         data=file, 
                    #         file_name=zip_file_name,
                    #         disabled=st.session_state.disabled)
                    #st.write(tmp_table)
                    #st.write(tmp_table)
                    if doc:
                        st.download_button(
                            label="Download Report",
                            data=bio.getvalue(),
                            file_name="Report.docx",
                            mime="docx",
                            disabled=st.session_state.disabled
                        )


                with col_d2:
        
                    # initiating a temp file
                    tmp_dir = tempfile.mkdtemp()

                    file_paths= []

                    for uploaded_file in st.session_state.pdf_files:
                        file_pth = os.path.join(tmp_dir, uploaded_file.name)
                        with open(file_pth, "wb") as file_opn:
                            file_opn.write(uploaded_file.getbuffer())
                        file_paths.append(file_pth)
#fetching file
                    for fetched_pdf in fetched_files:
                        # st.write(fetched_pdf)
                        file_pth = os.path.join('aml_docs/', fetched_pdf)
                        # st.write(file_pth)
                        file_paths.append(file_pth)

                
                    combined_doc_path = os.path.join(tmp_dir, "report.docx")
                    doc.save(combined_doc_path)



                    # Create a zip file with the uploaded PDF files and the combined document
                    zip_file_name = "package_files.zip"
                    if file_paths:
                        files =  [combined_doc_path] + file_paths
                        create_zip_file(files, zip_file_name)
                        # create_zip_file(file_paths, zip_file_name)
                    else:
                        pass

            
                    # Download the package

                    with open(zip_file_name, "rb") as file:
                        st.download_button(
                            label="Download Case Package", 
                            data=file, 
                            file_name=zip_file_name,
                            disabled=st.session_state.disabled)

                        # # Cleanup: Delete the temporary directory and its contents
                        # for file_path in file_paths + [combined_doc_path]:
                        #     os.remove(file_path)
                        # os.rmdir(temp_dir)

            with col6_up:   
                # Adding Radio button
                #st.markdown("""<span style="font-size: 24px; ">Make Decision</span>""", unsafe_allow_html=True)
                #if generate_button:
                    #text_data_doc = process_files_and_generate_responses(fetched_files)

                if st.session_state['llm'] == "Closed-Source":
                    st.markdown("""<span style="font-size: 24px;color:#0000FF">Is SAR filing required?</span>""", unsafe_allow_html=True)

                            
                    st.write("#### *SARA Recommendation*")
                    st.markdown("""<span style="font-size: 18px;">*Based on the following findings for the underlying case, under Bank Secrecy Act, it is recommended to file this case as a suspicious activity:*</span>""", unsafe_allow_html=True)
                    st.markdown("""<span style="font-size: 18px;">*1. A high-value transaction is made to a high-risk geography.*</span>""", unsafe_allow_html=True)
                    st.markdown("""<span style="font-size: 18px;">*2. There is an indication of suspicion with the involvement of multiple and frequent large cash deposits into Savings Account and corresponding debits through the Credit Card to a suspicious entity.*.</span>""", unsafe_allow_html=True)           
                
                    # query  = "Give your recommendation if SAR filling is required or not?"
                    # contexts = ', '.join(res_df_gpt['Answer'])
                    # prompt = f""" Summarize the context data provided with all the essential detials in it and also answer your recommendation on if SAR filling is required or not on the basis of summary?:
                    #     \n\n\
                    # Context: {contexts}\n\
                    # Question: {query}\n\
                    # Response: """
                    
            
                    
            
                    # response_sara_gpt = usellm(prompt) 
                    # response_sara_gpt = response_sara_gpt.replace("$", "USD ")
                    # #response_sara_gpt = response_sara_gpt.replace("10,000", "10,000 USD")
                    # #response_sara_gpt = response_sara_gpt.replace("10,600", "10,600 USD")
                    
                    # ##st.markdown(f'''<em>{response_sara_gpt}</em>''',unsafe_allow_html=True)
                    # st.markdown(f'''<em>{sara_close_source}</em>''',unsafe_allow_html=True)


                    st.warning('Please carefully review the recommendation and case details before the final submission',icon="⚠️")
                
                    # del(response_sara_gpt)
                
                elif st.session_state['llm'] == "Open-Source":
                    st.write("#### *SARA Recommendation*")
                    st.markdown("""<span style="font-size: 18px;">*Based on the following findings for the underlying case, under Bank Secrecy Act, it is recommended to file this case as a suspicious activity:*</span>""", unsafe_allow_html=True)
                    st.markdown("""<span style="font-size: 18px;">*1. A high-value transaction is made to a high-risk geography.*</span>""", unsafe_allow_html=True)
                    st.markdown("""<span style="font-size: 18px;">*2. There is an indication of suspicion with the involvement of multiple and frequent large cash deposits into Savings Account and corresponding debits through the Credit Card to a suspicious entity.*.</span>""", unsafe_allow_html=True)           
                
                    st.warning('Please carefully review the recommendation and case details before the final submission',icon="⚠️")         
    


                    # query  = "Give your recommendation if SAR filling is required or not?"
                    # contexts = ', '.join(res_df_llama['Answer'])
                    # prompt = f""" Summarize the context data provided with all the essential detials in it and also answer your recommendation on if SAR filling is required or not on the basis of summary?:
                    #     \n\n\
                    # Context: {contexts}\n\
                    # Question: {query}\n\
                    # Response: """
                    
                    
                    # response_sara_llama = llama_llm(llama_13b,prompt)
                    # response_sara_llama = response_sara_llama.replace("$", "USD ")
                    # # st.markdown(response1)
                    # st.markdown(f'''<em>{sara_open_source}</em>''',unsafe_allow_html=True)


                    # st.warning('Please carefully review the recommendation and case details before the final submission',icon="⚠️")
            
                    
                    
                    
            # st.markdown(
            #         """ <style>
            #                 div[role="radiogroup"] >  :first-child{
            #                     display: none !important;
            #                 }
            #             </style>
            #             """,
            #         unsafe_allow_html=True
            #     )
            # st.markdown("""<span style="font-size: 24px; ">Is SAR filing required?</span>""", unsafe_allow_html=True)
                selected_rad = st.radio(":blue", ["opt1", "Yes", "No", "Refer for review"], horizontal=True,disabled=st.session_state.disabled)
                if selected_rad == "Refer for review":
                    email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
                    email_id = st.text_input("Enter email ID")
                    if email_id and not re.match(email_regex, email_id):
                        st.error("Please enter a valid email ID")


                if st.button("Submit"):
                    if selected_rad in ("Yes"):
                        st.warning("Thanks for your review, your response has been submitted")
                    elif selected_rad in ("No"):
                        st.success("Thanks for your review, your response has been submitted")

                    else:
                        st.info("Thanks for your review, Case has been assigned to the next reviewer")


# Allow the user to clear all stored conversation sessions
                # if st.button("Reset Session"):
                #     reset_session_state()
                #     st.cache_data.clear()
                # #     pdf_files.clear()              



  
 
        

# Footer
st.markdown(
    """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    </style>
    """
    , unsafe_allow_html=True)
st.markdown('<div class="footer"><p></p></div>', unsafe_allow_html=True)

hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)


padding = 0
st.markdown(f""" <style>
    .reportview-container .main .block-container{{
        padding-top: {padding}rem;
        padding-right: {padding}rem;
        padding-left: {padding}rem;
        padding-bottom: {padding}rem;
    }} </style> """, unsafe_allow_html=True)



